from __future__ import annotations

from datetime import datetime
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent
GEN_DATE = "21 May 2026"

NAVY = "#0B2545"
TEAL = "#14B8A6"
BLUE = "#2563EB"
INDIGO = "#4F46E5"
GREEN = "#16A34A"
AMBER = "#D97706"
RED = "#DC2626"
SLATE = "#334155"
LIGHT = "#F8FAFC"
BORDER = "#CBD5E1"
MUTED = "#64748B"
WHITE = "#FFFFFF"


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


F_TITLE = font(46, True)
F_H1 = font(30, True)
F_H2 = font(22, True)
F_BODY = font(19)
F_SMALL = font(16)
F_TINY = font(14)


def hex_to_rgb(value: str):
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy, max_width: int, fnt, fill=SLATE, spacing=5, anchor=None):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        probe = f"{line} {word}".strip()
        if draw.textbbox((0, 0), probe, font=fnt)[2] <= max_width:
            line = probe
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    x, y = xy
    for idx, line in enumerate(lines):
        draw.text((x, y), line, font=fnt, fill=hex_to_rgb(fill), anchor=anchor)
        y += fnt.size + spacing
    return y


def rounded(draw, box, fill, outline=BORDER, width=2, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=width)


def arrow(draw, p1, p2, fill=SLATE, width=4):
    draw.line([p1, p2], fill=hex_to_rgb(fill), width=width)
    x1, y1 = p1
    x2, y2 = p2
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - direction * 14, y2 - 9), (x2 - direction * 14, y2 + 9)]
    else:
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - direction * 14), (x2 + 9, y2 - direction * 14)]
    draw.polygon(pts, fill=hex_to_rgb(fill))


def draw_icon(draw: ImageDraw.ImageDraw, center, kind: str, color: str, size: int = 42):
    cx, cy = center
    r = size // 2
    draw.ellipse((cx - r, cy - r, cx + r, cy + r), fill=hex_to_rgb(color))
    white = hex_to_rgb(WHITE)
    if kind == "human":
        head_r = max(5, size // 7)
        draw.ellipse((cx - head_r, cy - r + 8, cx + head_r, cy - r + 8 + head_r * 2), fill=white)
        body = (cx - size // 5, cy - r + 8 + head_r * 2 + 4, cx + size // 5, cy + r - 8)
        draw.rounded_rectangle(body, radius=max(4, size // 12), fill=white)
        draw.line((cx - size // 5, cy + 2, cx - size // 3, cy + size // 6), fill=white, width=max(2, size // 16))
        draw.line((cx + size // 5, cy + 2, cx + size // 3, cy + size // 6), fill=white, width=max(2, size // 16))
    elif kind == "agent":
        head = (cx - size // 4, cy - size // 5, cx + size // 4, cy + size // 5)
        draw.rounded_rectangle(head, radius=max(4, size // 12), outline=white, width=max(2, size // 16))
        eye_r = max(2, size // 18)
        draw.ellipse((cx - size // 8 - eye_r, cy - eye_r, cx - size // 8 + eye_r, cy + eye_r), fill=white)
        draw.ellipse((cx + size // 8 - eye_r, cy - eye_r, cx + size // 8 + eye_r, cy + eye_r), fill=white)
        draw.line((cx, cy - size // 5, cx, cy - size // 3), fill=white, width=max(2, size // 18))
        draw.ellipse((cx - 3, cy - size // 3 - 5, cx + 3, cy - size // 3 + 1), fill=white)
        draw.line((cx - size // 5, cy + size // 5, cx + size // 5, cy + size // 5), fill=white, width=max(2, size // 18))
    elif kind == "approval":
        draw.line((cx, cy - size // 4, cx, cy + size // 10), fill=white, width=max(3, size // 12))
        dot = max(3, size // 14)
        draw.ellipse((cx - dot, cy + size // 4 - dot, cx + dot, cy + size // 4 + dot), fill=white)
    elif kind == "assist":
        draw_icon(draw, center, "agent", color, size)
        check_w = max(3, size // 14)
        draw.line((cx + size // 10, cy + size // 5, cx + size // 5, cy + size // 3), fill=white, width=check_w)
        draw.line((cx + size // 5, cy + size // 3, cx + size // 3, cy + size // 10), fill=white, width=check_w)
    else:
        gear_r = size // 5
        draw.ellipse((cx - gear_r, cy - gear_r, cx + gear_r, cy + gear_r), outline=white, width=max(2, size // 16))
        for dx, dy in [(0, -1), (1, 0), (0, 1), (-1, 0)]:
            draw.line((cx + dx * gear_r, cy + dy * gear_r, cx + dx * (gear_r + 8), cy + dy * (gear_r + 8)), fill=white, width=max(2, size // 18))


def title_block(draw, title: str, subtitle: str, width: int):
    draw.text((70, 44), title, font=F_TITLE, fill=hex_to_rgb(NAVY))
    draw.text((72, 100), subtitle, font=F_BODY, fill=hex_to_rgb(MUTED))
    draw.rounded_rectangle((70, 136, width - 70, 142), radius=3, fill=hex_to_rgb(TEAL))


def save_jpg(img: Image.Image, name: str):
    path = OUT / name
    img.convert("RGB").save(path, "JPEG", quality=95, optimize=True)
    return path


def business_canvas():
    w, h = 2400, 1600
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(d, "Business Canvas", "Company AP automation platform - consolidated finance workflow", w)
    cols = [70, 530, 990, 1450, 1910, 2330]
    y0, y1, y2 = 190, 1110, 1480
    headers = [
        ("Key Partners", "Departments; vendors; AP finance; CFO; Xero; bank portal; legacy Google/Excel trackers; Trello stakeholders."),
        ("Key Activities", "Invoice capture; PO sync; agent verification; approval routing; AP review; WHT/voucher; Xero bill; bank execution; closeout."),
        ("Value Proposition", "One source of truth that cuts AP cycle time from about 3 days toward 1 day, with stronger controls and audit history."),
        ("Relationships", "Role-based portals, notifications, comments, missing-doc loops, approval/rejection reasons, immutable close state."),
        ("Customer Segments", "Department requesters, department heads, AP clerks, company admin, CFO, finance leadership."),
    ]
    colors = ["#E0F2FE", "#ECFDF5", "#EEF2FF", "#FFF7ED", "#F1F5F9"]
    for i, (head, body) in enumerate(headers):
        x0, x1 = cols[i], cols[i + 1] - 20
        rounded(d, (x0, y0, x1, y1), colors[i], outline=BORDER, radius=22)
        d.text((x0 + 28, y0 + 26), head, font=F_H1, fill=hex_to_rgb(NAVY))
        draw_wrapped(d, body, (x0 + 28, y0 + 80), x1 - x0 - 56, F_BODY)
        if head == "Value Proposition":
            rounded(d, (x0 + 36, y0 + 340, x1 - 36, y0 + 680), WHITE, outline=TEAL, radius=18)
            d.text((x0 + 62, y0 + 372), "Core outcomes", font=F_H2, fill=hex_to_rgb(NAVY))
            for j, item in enumerate(["Fewer duplicate entries", "Clear approval ownership", "Real bookkeeping trail", "Automated bank/Xero close path"]):
                d.ellipse((x0 + 62, y0 + 424 + j * 48, x0 + 78, y0 + 440 + j * 48), fill=hex_to_rgb(TEAL))
                d.text((x0 + 92, y0 + 414 + j * 48), item, font=F_BODY, fill=hex_to_rgb(SLATE))

    lower = [
        ("Key Resources", "React role-based UI, NestJS API, Prisma/Postgres, file attachments, ticket activity log, payment batch engine, Xero OAuth credentials."),
        ("Channels", "AP board, department invoice upload, head board, ticket detail, operations dashboard, Xero API, bank/payment gateway flow."),
        ("Cost Structure", "Implementation, hosting, Xero and bank API setup, user training, data migration, support, security/compliance maintenance."),
        ("Benefit Metrics", "Cycle time, first-pass approval rate, missing-doc count, bank failures, Xero sync errors, overdue tickets, manual touchpoints removed."),
    ]
    widths = [(70, 650), (680, 1260), (1290, 1810), (1840, 2330)]
    for i, (head, body) in enumerate(lower):
        x0, x1 = widths[i]
        rounded(d, (x0, y1 + 30, x1, y2), WHITE, outline=BORDER, radius=18)
        d.text((x0 + 24, y1 + 56), head, font=F_H2, fill=hex_to_rgb(NAVY))
        draw_wrapped(d, body, (x0 + 24, y1 + 104), x1 - x0 - 48, F_SMALL)
    return save_jpg(img, "01_business_canvas.jpg")


def process_flow():
    w, h = 2600, 1700
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(d, "End-to-End Process Flow", "Invoice to payment complete, including approval, exceptions, Xero, and payment gateway execution", w)
    columns = [
        ("Department intake", BLUE, [
            ("1", "Department uploads invoice and completes PO fields"),
            ("2", "System creates invoice, synced PO and AP ticket"),
            ("3", "Agent verifies required fields, vendor, PO and account"),
        ]),
        ("Approval gate", AMBER, [
            ("4", "Department head approves or rejects with reason"),
            ("5", "If rejected, department fixes and resubmits"),
        ]),
        ("AP finance processing", INDIGO, [
            ("6", "AP reviews documents and missing-doc loop"),
            ("7", "Vendor, PO and account verification"),
            ("8", "WHT filer/non-filer calculation and voucher"),
            ("9", "Xero AP bill is created"),
        ]),
        ("Payment and close", GREEN, [
            ("10", "AP prepares/uploads bank payment"),
            ("11", "CFO verifies and signs bank payment"),
            ("12", "Bank/payment gateway executes payment"),
            ("13", "Xero paid, requester notified, ticket locked"),
        ]),
    ]
    col_w, gap = 570, 44
    card_h = 116
    col_boxes = []
    for i, (title, color, items) in enumerate(columns):
        x = 70 + i * (col_w + gap)
        rounded(d, (x, 220, x + col_w, 1210), "#FFFFFF", outline=color, width=4, radius=22)
        d.text((x + 26, 252), title, font=F_H2, fill=hex_to_rgb(color))
        boxes = []
        y = 330
        for num, label in items:
            rounded(d, (x + 28, y, x + col_w - 28, y + card_h), "#F8FAFC", outline=BORDER, width=2, radius=16)
            d.ellipse((x + 50, y + 30, x + 88, y + 68), fill=hex_to_rgb(color))
            d.text((x + 62, y + 36), num, font=F_SMALL, fill=hex_to_rgb(WHITE))
            draw_wrapped(d, label, (x + 108, y + 24), col_w - 165, F_SMALL, fill=SLATE)
            boxes.append((x + 28, y, x + col_w - 28, y + card_h))
            y += card_h + 44
        col_boxes.append(boxes)
    for i in range(len(col_boxes) - 1):
        a = col_boxes[i][-1]
        b = col_boxes[i + 1][0]
        y = (a[1] + a[3]) // 2
        arrow(d, (a[2] + 4, y), (b[0] - 12, y), fill=SLATE, width=4)
    rounded(d, (110, 1300, 1230, 1535), "#FEF2F2", outline=RED, width=3, radius=20)
    d.text((145, 1332), "Exception loop", font=F_H2, fill=hex_to_rgb(RED))
    draw_wrapped(d, "Reject or missing documents sends the ticket back to department draft/rework with the reason. Department updates fields, comments or attachments, then resubmits.", (145, 1380), 1040, F_SMALL)
    rounded(d, (1370, 1300, 2490, 1535), "#ECFDF5", outline=GREEN, width=3, radius=20)
    d.text((1405, 1332), "Close controls", font=F_H2, fill=hex_to_rgb(GREEN))
    draw_wrapped(d, "PAYMENT_COMPLETE locks ticket fields for audit. Remaining or follow-up payments use a linked child ticket referencing the old record.", (1405, 1380), 1040, F_SMALL)
    return save_jpg(img, "02_process_flow.jpg")


def architecture_diagram():
    w, h = 2500, 1600
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(d, "Functional Architecture", "Role-based AP automation application architecture", w)
    layers = [
        ("Presentation Layer", "#DBEAFE", ["Login/Register", "AP Kanban Board", "Head Board", "Invoice Detail", "Ticket Detail", "Dashboard", "Operations"]),
        ("Application/API Layer", "#ECFDF5", ["Auth/JWT + RolesGuard", "Invoices Module", "Tickets Module", "Approvals Module", "Full Scope Ops", "Payments Module", "Departments/Vendors"]),
        ("Domain Services", "#EEF2FF", ["Due date rule: time cutoff", "WHT calculation", "PO/invoice sync", "Approval/rejection loop", "Ticket activity/comments", "Payment gateway", "Xero sync"]),
        ("Data and Storage", "#FFF7ED", ["Postgres via Prisma", "PaymentTicket", "Invoice/PO", "Vendor/User/Department", "SupportingDocument", "PaymentBatch/Record", "AuditLog/Notification"]),
        ("External Systems", "#F1F5F9", ["Xero Accounting API", "Meezan bank CSV", "Payment gateway", "Google/Excel migration", "Trello references", "File store"]),
    ]
    x0, x1 = 90, w - 90
    y = 210
    for title, fill, items in layers:
        rounded(d, (x0, y, x1, y + 210), fill, outline=BORDER, radius=22)
        d.text((x0 + 34, y + 26), title, font=F_H1, fill=hex_to_rgb(NAVY))
        card_w = (x1 - x0 - 90) // len(items)
        for i, item in enumerate(items):
            cx = x0 + 34 + i * (card_w + 10)
            rounded(d, (cx, y + 82, cx + card_w, y + 170), WHITE, outline=BORDER, radius=12)
            draw_wrapped(d, item, (cx + 14, y + 104), card_w - 28, F_TINY, fill=SLATE, spacing=3)
        y += 250
    for yy in [420, 670, 920, 1170]:
        arrow(d, (w // 2, yy), (w // 2, yy + 58), fill=TEAL, width=5)
    return save_jpg(img, "03_functional_architecture.jpg")


def integration_diagram():
    w, h = 2400, 1500
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(d, "Integration Diagram", "Implemented and planned integration points around the AP platform", w)
    center = (930, 690, 1470, 980)
    rounded(d, center, "#E0F2FE", outline=BLUE, width=5, radius=28)
    d.text((center[0] + 70, center[1] + 60), "AP Automation Core", font=F_H1, fill=hex_to_rgb(NAVY))
    draw_wrapped(d, "React web app + NestJS API + Prisma domain model", (center[0] + 70, center[1] + 116), 390, F_BODY)
    nodes = [
        ((100, 220, 650, 420), "Departments", "Invoice upload, PO fields, comments, attachments", BLUE),
        ((100, 610, 650, 810), "Legacy Sources", "Old Google Sheets, Excel trackers, prior Trello references", SLATE),
        ((100, 1000, 650, 1200), "Vendor Master", "Recurring/one-off vendors, bank account data, filer status inputs", GREEN),
        ((1750, 220, 2300, 420), "Xero", "OAuth, AP bill creation, payment marking, tenant status", GREEN),
        ((1750, 610, 2300, 810), "Bank / Payment Gateway", "Meezan CSV export, response import, payment gateway execution", INDIGO),
        ((1750, 1000, 2300, 1200), "Notifications", "Requester updates, approval comments, audit activities", AMBER),
        ((900, 1220, 1500, 1400), "File Storage", "Invoice scans, PO copies, bank confirmations, vouchers", TEAL),
    ]
    for box, title, body, color in nodes:
        rounded(d, box, WHITE, outline=color, width=4, radius=20)
        d.text((box[0] + 24, box[1] + 24), title, font=F_H2, fill=hex_to_rgb(color))
        draw_wrapped(d, body, (box[0] + 24, box[1] + 70), box[2] - box[0] - 48, F_SMALL)
        start = (box[2], (box[1] + box[3]) // 2) if box[0] < center[0] else (box[0], (box[1] + box[3]) // 2)
        end = (center[0], (center[1] + center[3]) // 2) if box[0] < center[0] else (center[2], (center[1] + center[3]) // 2)
        if box[1] > 1100:
            start = ((box[0] + box[2]) // 2, box[1])
            end = ((center[0] + center[2]) // 2, center[3])
        arrow(d, start, end, fill=color, width=4)
    rounded(d, (740, 220, 1660, 410), "#FEFCE8", outline=AMBER, width=3, radius=18)
    d.text((770, 250), "Integration status", font=F_H2, fill=hex_to_rgb(NAVY))
    draw_wrapped(d, "Xero contains real OAuth/API code when credentials are configured. Payment execution is modeled through bank batch flow and payment gateway integration; production bank/API response should drive final close.", (770, 292), 840, F_SMALL)
    return save_jpg(img, "04_integration_diagram.jpg")


def swimlane():
    w, h = 3000, 1900
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(d, "Swimlane Diagram", "Role ownership from invoice creation to payment complete", w)
    lanes = [
        ("Department User", "#DBEAFE"),
        ("System / Agent", "#ECFDF5"),
        ("AP Finance", "#E0F2FE"),
        ("CFO", "#FDE68A"),
        ("Bank / Payment Gateway", "#EEF2FF"),
        ("Xero", "#DCFCE7"),
        ("Requester", "#F1F5F9"),
    ]
    left, top, lane_h = 70, 210, 190
    for idx, (lane, fill) in enumerate(lanes):
        y = top + idx * lane_h
        rounded(d, (left, y, w - 70, y + lane_h - 12), fill, outline=BORDER, radius=14)
        d.text((left + 24, y + 26), lane, font=F_H2, fill=hex_to_rgb(NAVY))
    timeline = [
        "Intake", "System check", "AP review", "Data check",
        "Tax/voucher", "Xero bill", "Bank upload", "CFO sign",
        "Bank execute", "Xero paid", "Notify", "Close",
    ]
    start_x, step_gap = 350, 195
    for idx, label in enumerate(timeline):
        x = start_x + idx * step_gap
        d.text((x, 176), str(idx + 1), font=F_SMALL, fill=hex_to_rgb(TEAL))
        d.text((x + 24, 176), label, font=F_TINY, fill=hex_to_rgb(MUTED))
        d.line((x + 5, top - 10, x + 5, top + len(lanes) * lane_h - 24), fill=hex_to_rgb("#E2E8F0"), width=2)
    events = [
        (1, 0, "Create invoice\nand synced PO", BLUE),
        (2, 1, "Create ticket\nand run checks", TEAL),
        (3, 2, "Docs review\nand exceptions", BLUE),
        (4, 2, "Vendor/PO/account\nverification", BLUE),
        (5, 2, "WHT + voucher", BLUE),
        (6, 5, "Create Xero\nAP bill", GREEN),
        (7, 2, "Upload bank\npayment", BLUE),
        (8, 3, "CFO signs", AMBER),
        (9, 4, "Execute payment\nvia gateway", INDIGO),
        (10, 5, "Mark paid", GREEN),
        (11, 6, "Notify requester", SLATE),
        (12, 1, "Lock ticket\ncomplete", TEAL),
    ]
    card_w, card_h = 170, 82
    for step, lane, label, color in events:
        x = start_x + (step - 1) * step_gap - 32
        y = top + lane * lane_h + 62
        box = (x, y, x + card_w, y + card_h)
        rounded(d, box, WHITE, outline=color, width=3, radius=14)
        d.ellipse((x + 12, y + 16, x + 42, y + 46), fill=hex_to_rgb(color))
        d.text((x + 21, y + 20), str(step), font=F_TINY, fill=hex_to_rgb(WHITE))
        draw_wrapped(d, label.replace("\n", " "), (x + 50, y + 14), card_w - 60, F_TINY, fill=SLATE, spacing=2)
    rounded(d, (620, 350, 1280, 430), "#FEF2F2", outline=RED, width=2, radius=14)
    draw_wrapped(d, "Exception path: missing documents or AP rejection returns request to Department User for rework, then the flow restarts at AP review.", (650, 372), 600, F_TINY, fill=RED)
    return save_jpg(img, "05_swimlane_diagram.jpg")


def app_design_diagram():
    w, h = 2200, 1500
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(d, "Application Design", "Primary screens and role-scoped experience", w)
    screens = [
        ("Login/Register", "Demo + JWT auth, department selection, role registration", "#E0F2FE"),
        ("AP / Head Board", "Kanban columns, searchable tickets, role-permitted transitions", "#ECFDF5"),
        ("Invoice Detail", "Department invoice fields, PO sync, submit to head, approval status", "#EEF2FF"),
        ("Ticket Detail", "Workflow fields, comments, attachments, AP/CFO/payment gateway actions", "#FFF7ED"),
        ("Operations", "POs, approval matrix, payment batches, tax/GL reference, Xero status", "#F1F5F9"),
        ("Dashboard", "Metrics and role-scoped operational overview", "#F0FDFA"),
    ]
    x_positions = [110, 780, 1450]
    y_positions = [240, 670]
    for i, (title, body, fill) in enumerate(screens):
        x = x_positions[i % 3]
        y = y_positions[i // 3]
        rounded(d, (x, y, x + 560, y + 300), fill, outline=BORDER, radius=24)
        d.text((x + 30, y + 30), title, font=F_H1, fill=hex_to_rgb(NAVY))
        draw_wrapped(d, body, (x + 30, y + 90), 500, F_BODY)
    arrow(d, (670, 390), (780, 390), fill=TEAL, width=4)
    arrow(d, (1340, 390), (1450, 390), fill=TEAL, width=4)
    arrow(d, (1730, 540), (1730, 670), fill=TEAL, width=4)
    arrow(d, (1450, 820), (1340, 820), fill=TEAL, width=4)
    arrow(d, (780, 820), (670, 820), fill=TEAL, width=4)
    rounded(d, (360, 1090, 1840, 1335), WHITE, outline=BLUE, width=3, radius=20)
    d.text((400, 1125), "Design principles", font=F_H2, fill=hex_to_rgb(BLUE))
    draw_wrapped(d, "No manual ticket creation from AP; invoices generate board tickets. Every action is scoped by role, status, and field permissions. Closed tickets are read-only for audit. Comments and attachments preserve operational context.", (400, 1172), 1370, F_BODY)
    return save_jpg(img, "06_application_design.jpg")


def old_vs_new_artifacts_diagram():
    w, h = 2600, 1600
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(d, "Old vs New AP Artifacts", "Fragmented manual sources compared with the centralized AP automation system", w)

    left = (70, 210, 1170, 1390)
    right = (1430, 210, 2530, 1390)
    mid = (1210, 510, 1390, 1090)
    rounded(d, left, "#FFF7ED", outline=AMBER, width=4, radius=28)
    rounded(d, right, "#ECFDF5", outline=GREEN, width=4, radius=28)
    rounded(d, mid, "#EEF2FF", outline=INDIGO, width=3, radius=26)

    d.text((left[0] + 36, left[1] + 32), "Old Process: Manual + Fragmented", font=F_H1, fill=hex_to_rgb(AMBER))
    d.text((right[0] + 36, right[1] + 32), "New Process: Centralized + Controlled", font=F_H1, fill=hex_to_rgb(GREEN))
    d.text((mid[0] + 26, mid[1] + 40), "Transformation", font=F_H2, fill=hex_to_rgb(INDIGO))
    draw_wrapped(d, "Multiple sources and duplicate records move into one workflow, one board, one audit trail.", (mid[0] + 26, mid[1] + 88), 128, F_SMALL, fill=SLATE, spacing=4)
    arrow(d, (1170, 800), (1210, 800), fill=INDIGO, width=5)
    arrow(d, (1390, 800), (1430, 800), fill=INDIGO, width=5)

    old_sources = [
        ("Invoice sources", "Email invoices, repair/maintenance slips, Urdu payment notes, department-created invoice numbers"),
        ("Tracking tools", "Trello cards, Google Sheet v1, Google Sheet v2, Excel trackers, old reference sheets"),
        ("Finance checks", "Vendor account Excel lookup, PO validation, manual WHT filer/non-filer calculation"),
        ("Payment tools", "Manual voucher, bank portal upload, CFO bank sign, Xero paid marking"),
    ]
    issues = [
        "Duplicate entries across Trello, sheets, Excel and Xero",
        "No single owner for current status or blocker",
        "Missing documents discovered late",
        "Invoice number not always unique or available",
        "Vendor account mismatch risk",
        "Manual WHT and voucher error risk",
        "Cycle time around 3 days when follow-ups are manual",
    ]

    new_artifacts = [
        ("Single intake", "Department invoice upload, synced PO, mandatory invoice/PO fields, attachments"),
        ("Approval control", "Agent checks, department head read-only approval/rejection with reason and comments"),
        ("AP board", "Kanban status engine, assignments, due-date time cutoff, missing-doc loop, audit activity"),
        ("Finance execution", "Vendor/PO/account check, WHT + voucher, Xero bill, payment gateway/bank execution, paid close"),
    ]
    benefits = [
        "One source of truth for every AP ticket",
        "Invoice, PO, comments, attachments and approvals linked together",
        "Role-based access for department, head, AP, CFO and admin",
        "Xero bookkeeping trail and payment gateway/bank confirmation path",
        "Closed tickets become read-only for audit",
        "Target cycle moves from about 3 days toward 1 day",
    ]

    def stack_cards(origin_x, origin_y, items, color):
        y = origin_y
        for title, body in items:
            rounded(d, (origin_x, y, origin_x + 475, y + 135), WHITE, outline=color, width=2, radius=16)
            d.text((origin_x + 22, y + 18), title, font=F_H2, fill=hex_to_rgb(color))
            draw_wrapped(d, body, (origin_x + 22, y + 56), 420, F_TINY, fill=SLATE, spacing=2)
            y += 158

    stack_cards(left[0] + 36, left[1] + 110, old_sources, AMBER)
    stack_cards(right[0] + 36, right[1] + 110, new_artifacts, GREEN)

    rounded(d, (left[0] + 545, left[1] + 110, left[2] - 36, left[3] - 36), "#FEF2F2", outline=RED, width=3, radius=18)
    d.text((left[0] + 575, left[1] + 142), "Manual issues", font=F_H2, fill=hex_to_rgb(RED))
    y = left[1] + 195
    for item in issues:
        d.ellipse((left[0] + 575, y + 8, left[0] + 590, y + 23), fill=hex_to_rgb(RED))
        y = draw_wrapped(d, item, (left[0] + 608, y), 450, F_SMALL, fill=SLATE, spacing=4) + 10

    rounded(d, (right[0] + 545, right[1] + 110, right[2] - 36, right[3] - 36), "#F0FDFA", outline=TEAL, width=3, radius=18)
    d.text((right[0] + 575, right[1] + 142), "Centralized outcomes", font=F_H2, fill=hex_to_rgb(TEAL))
    y = right[1] + 195
    for item in benefits:
        d.ellipse((right[0] + 575, y + 8, right[0] + 590, y + 23), fill=hex_to_rgb(TEAL))
        y = draw_wrapped(d, item, (right[0] + 608, y), 450, F_SMALL, fill=SLATE, spacing=4) + 10

    rounded(d, (400, 1450, 2200, 1530), WHITE, outline=BLUE, width=3, radius=20)
    d.text((440, 1474), "Message for stakeholders:", font=F_H2, fill=hex_to_rgb(BLUE))
    draw_wrapped(d, "Before: invoices and payment records were scattered across people and tools. Now: every request becomes one controlled AP ticket connected to invoice, PO, Xero, payment gateway, comments, attachments and audit history.", (720, 1476), 1420, F_SMALL, fill=SLATE, spacing=3)
    return save_jpg(img, "07_old_vs_new_artifacts.jpg")


def human_agent_swimlane_diagram():
    w, h = 3200, 2050
    img = Image.new("RGB", (w, h), hex_to_rgb(LIGHT))
    d = ImageDraw.Draw(img)
    title_block(
        d,
        "Human + AI Agent Swimlane",
        "Clear ownership of manual approvals, AI automation, semi-autonomous finance work, and monitoring",
        w,
    )

    legend_y = 158
    legend = [
        ("Human action", AMBER),
        ("AI agent automation", TEAL),
        ("Semi-autonomous assist", INDIGO),
        ("Mandatory approval", RED),
    ]
    x = 80
    for label, color in legend:
        rounded(d, (x, legend_y, x + 310, legend_y + 52), WHITE, outline=color, width=3, radius=16)
        icon_kind = {
            "Human action": "human",
            "AI agent automation": "agent",
            "Semi-autonomous assist": "assist",
            "Mandatory approval": "approval",
        }[label]
        draw_icon(d, (x + 34, legend_y + 26), icon_kind, color, size=30)
        d.text((x + 62, legend_y + 15), label, font=F_SMALL, fill=hex_to_rgb(SLATE))
        x += 340

    lanes = [
        ("Department User", "Human", "#FFF7ED", "human", AMBER),
        ("Intake / OCR Agent", "AI", "#ECFDF5", "agent", TEAL),
        ("Validation / Risk Agent", "AI", "#F0FDFA", "agent", TEAL),
        ("AP Clerk", "Human + assist", "#E0F2FE", "human", AMBER),
        ("Finance Agents", "AI assist", "#EEF2FF", "assist", INDIGO),
        ("CFO", "Human approval", "#FDE68A", "approval", RED),
        ("Payment Gateway / Bank Agent", "System + AI", "#F1F5F9", "agent", TEAL),
        ("Monitoring / Notification Agent", "AI", "#DCFCE7", "agent", TEAL),
    ]

    left, top, lane_h = 70, 285, 176
    lane_right = w - 70
    label_w = 315
    for idx, (lane, owner, fill, icon_kind, icon_color) in enumerate(lanes):
        y = top + idx * lane_h
        rounded(d, (left, y, lane_right, y + lane_h - 12), fill, outline=BORDER, radius=16)
        draw_icon(d, (left + 42, y + 48), icon_kind, icon_color, size=46)
        d.text((left + 78, y + 26), lane, font=F_H2, fill=hex_to_rgb(NAVY))
        d.text((left + 78, y + 68), owner, font=F_SMALL, fill=hex_to_rgb(MUTED))

    stages = [
        "Intake",
        "Extract",
        "Validate",
        "AP Review",
        "Tax / Xero",
        "CFO Sign",
        "Pay / Reconcile",
        "Close / Report",
    ]
    stage_x0 = left + label_w
    usable_w = lane_right - stage_x0 - 20
    stage_gap = usable_w // len(stages)
    for idx, stage in enumerate(stages):
        sx = stage_x0 + idx * stage_gap
        d.line((sx, top - 16, sx, top + lane_h * len(lanes) - 22), fill=hex_to_rgb("#CBD5E1"), width=2)
        d.text((sx + 10, top - 45), f"{idx + 1}", font=F_SMALL, fill=hex_to_rgb(TEAL))
        draw_wrapped(d, stage, (sx + 34, top - 48), stage_gap - 40, F_TINY, fill=MUTED, spacing=1)
    d.line((lane_right - 90, top - 16, lane_right - 90, top + lane_h * len(lanes) - 22), fill=hex_to_rgb("#CBD5E1"), width=2)

    def step_card(stage_idx, lane_idx, text, kind, height=92):
        colors = {
            "human": AMBER,
            "agent": TEAL,
            "assist": INDIGO,
            "approval": RED,
            "system": SLATE,
        }
        color = colors[kind]
        sx = stage_x0 + stage_idx * stage_gap + 18
        y = top + lane_idx * lane_h + 56
        box = (sx, y, sx + stage_gap - 36, y + height)
        fill = WHITE if kind != "approval" else "#FEF2F2"
        rounded(d, box, fill, outline=color, width=3, radius=14)
        icon_kind = {"human": "human", "agent": "agent", "assist": "assist", "approval": "approval", "system": "agent"}[kind]
        draw_icon(d, (sx + 30, y + 32), icon_kind, color, size=34)
        draw_wrapped(d, text, (sx + 58, y + 15), stage_gap - 112, F_TINY, fill=SLATE, spacing=2)
        return box

    events = [
        (0, 0, "Upload invoice / email / slip and add basic department info", "human"),
        (1, 1, "Read PDF/image/email, OCR, classify bill type, extract fields", "agent"),
        (2, 2, "Check duplicate, missing docs, PO sync, vendor/account mismatch, confidence score", "agent", 112),
        (3, 3, "Review exceptions, verify docs, decide if ready for finance processing", "human"),
        (3, 0, "Fix missing docs or AP rejection, add proof in comments/attachments", "human"),
        (4, 4, "Suggest WHT filer/non-filer, voucher draft, GL/expense coding, Xero bill payload", "assist", 112),
        (4, 3, "Accept/adjust AI suggestions and submit Xero/payment preparation", "human"),
        (5, 5, "CFO signs payment. Mandatory human approval.", "approval"),
        (6, 6, "Execute payment through gateway/bank, capture confirmation and status", "system"),
        (6, 4, "Reconcile bank confirmation with ticket and Xero payment record", "agent"),
        (7, 7, "Notify requester, update board, SLA reporting, anomaly alerts, close recommendation", "agent", 112),
        (7, 3, "Final exception handling only if reconciliation or close confidence fails", "human", 112),
    ]
    for item in events:
        step_card(*item)

    # Main happy-path arrows across stages.
    arrow_y = top + lane_h * len(lanes) + 10
    rounded(d, (left, arrow_y + 18, lane_right, arrow_y + 135), WHITE, outline=BLUE, width=3, radius=20)
    d.text((left + 36, arrow_y + 48), "Human-in-the-loop rule", font=F_H2, fill=hex_to_rgb(BLUE))
    draw_wrapped(
        d,
        "AI agents can extract, validate, recommend, notify, reconcile and monitor. AP exception review, CFO sign, vendor bank account changes, tax overrides, and high-value payment release must remain human-approved.",
        (left + 330, arrow_y + 50),
        lane_right - left - 390,
        F_SMALL,
        fill=SLATE,
        spacing=3,
    )
    return save_jpg(img, "08_human_agent_swimlane.jpg")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_text(cell, text: str, bold=False, color=SLATE):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document, title: str, subtitle: str):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p2 = doc.add_paragraph()
    r2 = p2.add_run(subtitle)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor.from_string("64748B")
    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Generated: {GEN_DATE} | Project: Company AP Automation")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor.from_string("64748B")


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_image(doc: Document, path: Path, caption: str):
    doc.add_picture(str(path), width=Inches(6.6))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string("64748B")


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=NAVY)
        set_cell_shading(hdr[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
    doc.add_paragraph()
    return table


def save_doc(doc: Document, path: Path):
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_Updated{path.suffix}")
        doc.save(fallback)
        return fallback


workflow_steps = [
    ("1", "Department creates invoice", "DEPT_USER / COMPANY_ADMIN", "Invoice + synced PO draft, AP ticket NEW_REQUEST"),
    ("2", "Agent verification", "System", "Required fields, vendor, account, PO and amount checks"),
    ("3", "Department head decision", "DEPT_ADMIN", "Approve to finance or reject with reason"),
    ("4", "AP document review", "AP_CLERK", "Complete docs or missing-doc loop to requester"),
    ("5", "Data verification", "AP_CLERK", "Vendor, PO, invoice account, old sheet/reference validation"),
    ("6", "WHT and voucher", "AP_CLERK", "Filer/non-filer rate, WHT amount, net payable, voucher"),
    ("7", "Xero bill", "AP_CLERK / COMPANY_ADMIN", "AP bill sync and Xero bill identifiers"),
    ("8", "Payment preparation", "AP_CLERK", "Payment file/batch and bank portal reference"),
    ("9", "CFO sign", "CFO", "CFO authorizes bank portal payment"),
    ("10", "Bank execution", "AP / Payment Gateway", "Executed status or payment gateway confirmation"),
    ("11", "Xero paid and close", "AP / System", "Mark paid, notify requester, lock ticket"),
]


roles = [
    ("DEPT_USER", "Create invoice/PO, complete department fields, respond to rejection or missing docs."),
    ("DEPT_ADMIN", "View department head board, read ticket, approve or reject with reason."),
    ("AP_CLERK", "Run finance workflow, assign AP/CFO, verify docs/data, WHT, Xero, bank and close."),
    ("CFO", "Open CFO sign pending tickets and record bank authorization."),
    ("COMPANY_ADMIN", "Full administrative scope across departments, operations, Xero setup and rules."),
]


def business_canvas_doc(diagrams):
    doc = Document()
    style_doc(doc, "Business Canvas", "AP automation value model and operating case")
    doc.add_heading("Executive View", level=1)
    doc.add_paragraph("The platform replaces a fragmented Trello, Google Sheets, Excel and manual bank/Xero process with a role-scoped AP workflow. It centralizes invoice creation, PO sync, department head approval, AP verification, CFO bank authorization, bank execution, Xero bookkeeping and payment closeout.")
    add_image(doc, diagrams["canvas"], "Business canvas for Company AP Automation")
    doc.add_heading("Canvas Detail", level=1)
    add_table(doc, ["Area", "Project Interpretation"], [
        ("Customer segments", "Departments, department heads, AP finance, CFO, company admin and finance leadership."),
        ("Value proposition", "Reduce duplicate entry and compress AP processing from roughly 3 days toward 1 day with better controls."),
        ("Channels", "Kanban board, invoice pages, head board, operations dashboard, Xero and bank/payment gateway flows."),
        ("Key activities", "Invoice capture, agent verification, approvals, document review, WHT, voucher, payment and reconciliation."),
        ("Key resources", "React UI, NestJS API, Prisma/Postgres schema, file attachments, role permissions and integrations."),
        ("Partners", "Xero, bank portal, vendors, Trello/legacy references and Google/Excel migration sources."),
        ("Metrics", "Cycle time, overdue tickets, first pass approval, missing-doc loop count, bank failures, Xero errors."),
    ])
    path = OUT / "AP_Automation_Business_Canvas.docx"
    return save_doc(doc, path)


def application_design_doc(diagrams):
    doc = Document()
    style_doc(doc, "Application Design", "Functional, UX and technical design for the AP automation platform")
    doc.add_heading("Design Summary", level=1)
    doc.add_paragraph("The app is a React and NestJS AP workflow system with role-based screens, status-scoped editing, a Kanban ticket board, invoice/PO detail pages, AP operations controls and external-system adapters.")
    add_image(doc, diagrams["app"], "Application design overview")
    doc.add_heading("Role-Based Experience", level=1)
    add_table(doc, ["Role", "Primary Experience"], roles)
    doc.add_heading("Core Screens", level=1)
    add_table(doc, ["Screen", "Purpose"], [
        ("Login/Register", "Authenticate users and bind department and role."),
        ("AP Board / Head Board", "Track tickets by status; department head sees only approval board."),
        ("Invoice Detail", "Complete invoice/PO fields, submit to head and show approval status."),
        ("Ticket Detail", "Run AP actions, comments, attachments, CFO sign, payment gateway close, Xero close."),
        ("Operations", "Purchase orders, queries, approval matrix, Meezan batches, tax/GL references and Xero status."),
        ("Dashboard", "Role-scoped metrics and workload summary."),
    ])
    doc.add_heading("Data Design", level=1)
    add_bullets(doc, [
        "Invoice and PurchaseOrder records are created/synced from department submission.",
        "PaymentTicket is the operational tracking record and carries status, assignment, vendor, PO, WHT, bank and Xero fields.",
        "SupportingDocument enables invoice scans, PO copies, vouchers and bank confirmations.",
        "TicketActivity records comments, status changes, attachments and payment gateway actions.",
        "PaymentBatch and PaymentRecord support Meezan export and bank response/reconciliation scope.",
    ])
    doc.add_heading("Controls", level=1)
    add_bullets(doc, [
        "Role permissions restrict which statuses each user can move.",
        "Field-level edit rules prevent AP, CFO, department and head users from changing out-of-scope data.",
        "Payment complete tickets are locked for audit.",
        "Due date is 3 days from the current day if created before the configured cutoff time, otherwise 3 days from the next day.",
    ])
    add_image(doc, diagrams["architecture"], "Functional architecture diagram")
    path = OUT / "AP_Automation_Application_Design.docx"
    return save_doc(doc, path)


def prd_doc(diagrams):
    doc = Document()
    style_doc(doc, "Product Requirements Document (PRD)", "Company AP Automation")
    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph("The current AP process uses Trello, Google Sheets, Excel sheets, email invoices, bank portal activity and Xero bookkeeping separately. This creates duplicate entry, missing status visibility, slower approvals, higher error risk and unclear ownership.")
    doc.add_heading("Goals", level=1)
    add_bullets(doc, [
        "Centralize AP requests into one workflow and board.",
        "Reduce manual touchpoints and target a one-day processing cycle where documents are complete.",
        "Make department head approval explicit before AP finance processing.",
        "Integrate bookkeeping through Xero and payment execution through bank/payment gateway workflow.",
        "Preserve audit trail through comments, attachments and locked completed tickets.",
    ])
    doc.add_heading("Out of Scope / Future Production Hardening", level=1)
    add_bullets(doc, [
        "Production bank API credentials and payment gateway callback/response handling.",
        "Full Trello API synchronization beyond reference fields.",
        "Automated OCR quality guarantees for every invoice format.",
    ])
    doc.add_heading("Personas", level=1)
    add_table(doc, ["Persona", "Need"], roles)
    doc.add_heading("Functional Requirements", level=1)
    add_table(doc, ["ID", "Requirement", "Acceptance Criteria"], [
        ("FR-01", "Department invoice creation", "Department user can create/upload invoice; system creates linked PO and AP ticket."),
        ("FR-02", "Agent verification", "Request cannot move to head until key invoice, PO, vendor and amount checks pass."),
        ("FR-03", "Head approval/rejection", "Head sees read-only request, can approve or reject with reason; buttons hide after decision."),
        ("FR-04", "AP workflow board", "AP sees finance statuses and can move only permitted transitions."),
        ("FR-05", "Missing document loop", "AP can mark incomplete, requester can fix and resubmit."),
        ("FR-06", "WHT and voucher", "Filer/non-filer rate calculates WHT and net payable; voucher number stored."),
        ("FR-07", "Xero bill and paid marking", "Xero bill and payment IDs are stored; failures are visible."),
        ("FR-08", "Bank execution", "CFO signs; AP can execute manually or receive payment gateway confirmation."),
        ("FR-09", "Attachments/comments", "Users can add/download allowed attachments and comments by role/stage."),
        ("FR-10", "Audit lock", "PAYMENT_COMPLETE ticket cannot be edited."),
    ])
    doc.add_heading("Workflow", level=1)
    add_table(doc, ["Step", "Stage", "Owner", "Output"], workflow_steps)
    add_image(doc, diagrams["process"], "End-to-end process flow")
    add_image(doc, diagrams["swimlane"], "Swimlane role ownership")
    doc.add_heading("Success Metrics", level=1)
    add_bullets(doc, [
        "Average invoice-to-paid cycle time.",
        "Percentage of tickets completed within due date.",
        "First pass department head approval rate.",
        "Missing-doc loop count by department/vendor.",
        "Xero sync error rate and bank execution failure rate.",
    ])
    path = OUT / "AP_Automation_PRD.docx"
    return save_doc(doc, path)


def agentic_pdr_doc(diagrams):
    doc = Document()
    style_doc(
        doc,
        "Agentic AI Product Design Requirements (PDR)",
        "Implementation blueprint for the Company AP Automation human + AI agent workflow",
    )

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This PDR defines the next implementation scope for adding Agentic AI capabilities to the AP Automation system. "
        "The objective is to reduce manual finance operations while keeping department head approval, CFO payment sign-off, "
        "vendor bank-account changes, tax overrides and high-value payment release under mandatory human control."
    )
    doc.add_heading("Scope Objectives", level=2)
    add_table(doc, ["Objective Area", "Detailed Objective", "Target Outcome"], [
        ("Business consolidation", "Replace Trello, Google Sheets, Excel trackers, email follow-ups and manual status chasing with one AP workflow.", "Single source of truth for every invoice, PO, approval, payment and closeout."),
        ("Cycle-time reduction", "Use agents to remove low-value manual checking, reminders, extraction and reconciliation work.", "Move clean invoice processing from roughly 3 days toward 1 working day."),
        ("Control and audit", "Keep high-risk financial decisions under human control while logging every system and agent action.", "Stronger audit trail without slowing down AP operations."),
        ("Agentic automation", "Introduce AI agents for extraction, validation, recommendation, notification, monitoring and workflow QA.", "Human teams focus on exceptions, approvals and judgment-based finance decisions."),
        ("Clickable demo readiness", "Define screens, flows, states and interactions clearly enough to create a full Figma clickable prototype.", "Stakeholders can test the future workflow before engineering implementation."),
    ])
    doc.add_heading("In Scope", level=2)
    add_bullets(doc, [
        "Department invoice upload, manual invoice entry and email-import-ready intake flow.",
        "Invoice and PO synchronized detail experience with extracted fields, confidence and corrections.",
        "AI validation for duplicate, missing document, PO sync, vendor match, account mismatch and risk score.",
        "Department head approval/rejection portal with AI summary but mandatory human decision.",
        "AP finance board with AI readiness badges, recommendations and exception handling.",
        "WHT/voucher recommendation, Xero bill assistant and payment gateway/bank reconciliation support.",
        "Notification, SLA monitoring, reporting, audit log and QA workflow agents.",
        "Figma clickable demo specification for all primary personas and critical paths.",
    ])
    doc.add_heading("Out of Scope for First Implementation", level=2)
    add_bullets(doc, [
        "Fully autonomous payment approval or payment release without CFO/human sign-off.",
        "Automatic approval of vendor bank account changes.",
        "Replacing Xero as the accounting system of record.",
        "Guaranteeing perfect OCR for every invoice/slip format without human review fallback.",
        "Full Trello two-way sync; legacy references can be retained/imported first.",
    ])
    add_image(doc, diagrams["human_agent"], "Target human + AI agent swimlane")

    doc.add_heading("2. Current System Baseline", level=1)
    add_bullets(doc, [
        "Frontend: React screens for login/register, upload, invoice detail, AP/head board, ticket detail, operations and dashboard.",
        "Backend: NestJS modules for auth, invoices, approvals, tickets, vendors, departments, payments and full-scope operations.",
        "Workflow: invoice/PO creation, department head approval, AP document review, vendor/PO/account verification, WHT, voucher, Xero, payment gateway/bank execution and close.",
        "Data model: Prisma models for Invoice, PurchaseOrder, PaymentTicket, SupportingDocument, TicketActivity, PaymentBatch, PaymentRecord, Notification and XeroConnection.",
        "Controls: role-based transitions, field-level edit restrictions, comments, attachments, downloadable documents and locked payment-complete tickets.",
    ])
    doc.add_heading("Stakeholders and Personas", level=2)
    add_table(doc, ["Persona", "Current Pain", "Future Agentic Experience"], [
        ("Department User", "Uploads invoices and manually responds to missing-doc or rejection feedback.", "Gets AI-guided required fields, missing-doc prompts, upload confidence and clear status visibility."),
        ("Department Head", "Needs quick context before approving but should not edit finance data.", "Sees AI summary, risk flags, attachments and reason box; approve/reject remains human-only."),
        ("AP Clerk", "Manually checks docs, PO, vendor account, WHT, voucher, Xero and payment statuses.", "Gets validation checklist, recommended next action, WHT/voucher draft and exception-focused work queue."),
        ("CFO", "Needs confidence before signing payment without scanning every operational detail.", "Sees payment risk, vendor/account verification, amount summary and mandatory sign action."),
        ("Company Admin", "Maintains roles, vendors, departments, rules and integrations.", "Configures thresholds, agent rules, integration health and audit exports."),
        ("Auditor/Finance Leadership", "Needs evidence of approvals, payments and control adherence.", "Uses immutable ticket timeline, agent logs, human approvals and reconciliation reports."),
    ])
    doc.add_heading("Business Rules to Preserve", level=2)
    add_bullets(doc, [
        "Department creates invoices; AP should not manually create separate tickets for invoices.",
        "Every invoice creates or syncs a PO record before finance processing.",
        "Department head approval must occur before AP finance workflow begins.",
        "Rejected or missing-doc tickets return to the department with reason/comments.",
        "Due date uses configured time cutoff logic: before cutoff, count from current day; after cutoff, count from next day.",
        "Payment complete tickets are read-only and remain available for audit/history.",
    ])

    doc.add_heading("3. Target Agentic Operating Model", level=1)
    doc.add_paragraph(
        "The workflow should evolve from a manually operated Kanban board into an event-driven agentic operating layer. "
        "Agents listen to workflow events, perform bounded analysis or automation, write recommendations and low-risk updates through the existing API, and escalate high-risk decisions to humans."
    )
    add_table(doc, ["Layer", "Responsibility", "Implementation Direction"], [
        ("Workflow engine", "Keeps ticket status, permissions and audit trail authoritative.", "Existing NestJS ticket/invoice services remain source of truth."),
        ("Agent orchestrator", "Runs agents when invoice/ticket/payment events occur.", "Add background worker with event queue, idempotent job table and retry policy."),
        ("AI agents", "Extract, validate, recommend, notify, reconcile and monitor.", "Agents call approved tools/APIs; outputs stored as recommendations and activity records."),
        ("Human approvals", "Approve/reject department requests and authorize payment.", "Department Head and CFO actions stay manual with explicit UI gates."),
    ])

    doc.add_heading("4. AI Agent Roles and Responsibilities", level=1)
    add_table(doc, ["Agent", "Trigger", "Responsibilities", "Output"], [
        ("Document Intake Agent", "Invoice upload or email import", "Classify file, OCR PDF/image/slip, extract vendor, amount, invoice date, account, PO and bill type.", "Extracted fields, confidence score, document classification."),
        ("Validation/Risk Agent", "Invoice draft created or edited", "Check duplicate invoice, missing docs, PO sync, vendor match, account mismatch and old sheet references.", "Validation checklist, risk score, missing-doc list."),
        ("Approval Routing Agent", "Ready for head approval", "Route to correct department head, calculate SLA, remind/escalate overdue approvals.", "Approval task, reminders, escalation activity."),
        ("AP Finance Copilot", "Ticket enters AP review", "Recommend next status, highlight blockers, summarize invoice/PO/vendor context.", "AP recommendation panel and suggested action."),
        ("WHT/Voucher Agent", "Ticket enters WHT/voucher stage", "Suggest filer/non-filer rate, WHT amount, net payable, voucher draft and GL/expense mapping.", "Tax/voucher draft requiring AP confirmation."),
        ("Xero Sync Agent", "Ready for Xero bill or paid marking", "Prepare bill payload, detect sync errors, refresh token, retry safe failures.", "Xero bill/payment status and error resolution guidance."),
        ("Payment Gateway Agent", "Payment prepared or bank confirmation received", "Validate payment batch data, capture payment confirmation, reconcile with ticket and Xero status.", "Payment execution status and reconciliation result."),
        ("Notification Agent", "Missing docs, approval pending, overdue, close", "Send requester/head/AP/CFO notifications, write comments, summarize blockers.", "Notifications and audit activity."),
        ("QA/Workflow Agent", "Scheduled or after deployment", "Run end-to-end test scenarios and validate role permissions/status transitions.", "Test report and failed workflow alerts."),
    ])

    doc.add_heading("5. Human-in-the-Loop Rules", level=1)
    add_table(doc, ["Decision", "Automation Level", "Reason"], [
        ("Department head approval/rejection", "Human mandatory", "Business owner must approve budget/department request."),
        ("CFO payment sign", "Human mandatory", "Financial authorization and fraud control."),
        ("Vendor bank account creation/change", "Human mandatory with dual control", "High fraud risk; AI can flag mismatch but not approve."),
        ("WHT/tax override", "Human mandatory", "Compliance and tax liability risk."),
        ("High-value payment release", "Human mandatory", "Payment risk threshold requires explicit sign-off."),
        ("OCR/extraction correction", "Semi-autonomous", "AI can prefill; user confirms low-confidence fields."),
        ("Missing-doc reminders", "Autonomous", "Low-risk operational communication."),
        ("Status close after confirmed reconciliation", "Semi-autonomous", "Agent can recommend close; auto-close only when confidence and rules pass."),
    ])

    doc.add_heading("6. Functional Requirements for Implementation", level=1)
    add_table(doc, ["ID", "Requirement", "Acceptance Criteria"], [
        ("AG-01", "Agent job orchestration", "System creates idempotent background jobs for upload, validation, notification, Xero and payment events."),
        ("AG-02", "AI extraction panel", "Invoice detail shows extracted fields, confidence, source document and user-confirmed corrections."),
        ("AG-03", "Validation checklist", "Ticket shows required docs, duplicate risk, vendor/account match, PO sync and readiness status."),
        ("AG-04", "Recommendation actions", "AP sees suggested next action but final status move follows role permissions."),
        ("AG-05", "Approval reminders", "Head/CFO pending tasks trigger reminders and escalation based on SLA."),
        ("AG-06", "WHT/voucher draft", "Agent calculates draft WHT/net payable and AP must confirm before voucher is finalized."),
        ("AG-07", "Xero assistant", "Agent prepares bill/payment payload, stores sync error explanations and prevents duplicate bill creation."),
        ("AG-08", "Payment gateway reconciliation", "Agent matches gateway/bank confirmation to ticket, amount, vendor and Xero payment status."),
        ("AG-09", "Conversation assistant", "Users can ask ticket-specific questions and receive answers grounded in ticket, invoice, PO and activity data."),
        ("AG-10", "Workflow QA agent", "Automated tests verify critical role/status paths before release."),
    ])

    doc.add_heading("7. Data Model Additions", level=1)
    add_table(doc, ["Entity / Field", "Purpose"], [
        ("AgentJob", "Track queued/running/completed/failed agent jobs with idempotency key and retry count."),
        ("AgentRun", "Store model/tool used, prompt version, input references, output summary, confidence and cost metadata."),
        ("AgentRecommendation", "Store suggested field updates, next status, risk flags and human acceptance/rejection."),
        ("Invoice.extractionConfidence", "Field-level confidence for extracted values."),
        ("PaymentTicket.riskScore", "Aggregated duplicate/vendor/account/payment risk score."),
        ("PaymentTicket.agentReadiness", "READY, NEEDS_REVIEW, BLOCKED or FAILED state for agent checks."),
        ("TicketActivity.actorType", "Differentiate human, system and AI agent actions in audit history."),
    ])

    doc.add_heading("8. APIs and Tools Required", level=1)
    add_bullets(doc, [
        "LLM with vision/OCR capability for invoice and slip extraction.",
        "Embedding/fuzzy-match service for duplicate detection and vendor matching.",
        "Xero Accounting API with idempotency protection for bill/payment creation.",
        "Payment gateway or bank file/API response integration for payment execution status.",
        "Email ingestion API for invoices arriving by email.",
        "Notification channels: in-app first; optional email, Slack, Teams or WhatsApp later.",
        "Policy/rules engine for mandatory fields, approval thresholds, WHT rates and risk gates.",
    ])

    doc.add_heading("9. User Experience Changes", level=1)
    add_bullets(doc, [
        "Invoice Detail: add AI extraction review card with confidence indicators and confirm/correct buttons.",
        "Ticket Detail: add AI validation checklist, recommendation summary and risk reasons.",
        "AP Board: add agent-ready, blocked-by-docs and high-risk badges.",
        "Department Head Board: show AI summary but keep approve/reject buttons human-only.",
        "CFO View: show payment risk summary, vendor/account confirmation and payment gateway status.",
        "Operations: add Agent Jobs, failed runs, prompt/version settings and audit export.",
    ])
    doc.add_heading("Figma Clickable Demo Scope", level=2)
    add_table(doc, ["Screen", "Demo Purpose", "Must-Have Interactions"], [
        ("Login / Role Selector", "Show how each persona enters a role-scoped experience.", "Select demo user, login, route to correct dashboard/board."),
        ("Department Intake", "Create invoice, upload file, see AI extraction and synced PO fields.", "Upload document, review extracted fields, correct low-confidence field, submit to head."),
        ("Invoice Detail", "Show invoice and PO fields side by side with AI validation.", "Edit fields, view confidence, add attachment, submit approval, see status change."),
        ("Department Head Board", "Review pending approvals only.", "Open card, read AI summary, view docs, approve or reject with reason."),
        ("AP Board", "Show operational Kanban with human/agent ownership.", "Filter by priority/risk, open ticket, move only permitted statuses."),
        ("Ticket Detail", "Main AP workspace for validation, comments, attachments, WHT, Xero and payment.", "Accept/reject recommendation, add comment, download attachment, run Xero/payment actions."),
        ("CFO Sign View", "Show mandatory approval gate.", "Review payment summary, risk flags, sign payment, send to execution pending."),
        ("Operations / Agent Monitor", "Admin view of integrations, agent jobs and failed runs.", "Inspect failed agent job, retry safe job, open Xero/payment gateway status."),
        ("Dashboard / Reports", "Executive view of cycle time, overdue, agent savings and risk.", "Switch metrics, view drill-down, export report placeholder."),
        ("Conversational Assistant", "Ticket-aware chat for users.", "Ask what is missing, why blocked, next step, or summarize ticket."),
    ])
    doc.add_heading("Clickable Prototype Paths", level=2)
    add_table(doc, ["Path", "Start", "End State"], [
        ("Happy path", "Department user uploads clean invoice.", "Ticket payment complete with requester notified and audit timeline."),
        ("Rejected by head", "Department user submits incomplete/incorrect invoice.", "Head rejects with reason; department fixes and resubmits."),
        ("Missing docs loop", "AP finds missing PO/supporting document.", "Notification sent; requester uploads missing doc; AP resumes review."),
        ("High-risk vendor account", "Validation Agent detects account mismatch.", "Ticket is blocked until AP/admin human verification."),
        ("CFO payment approval", "AP prepares payment.", "CFO signs; payment gateway receives execution request."),
        ("Xero failure", "Xero sync returns error.", "Xero Sync Agent explains failure and AP retries after correction."),
    ])
    doc.add_heading("Design System Requirements for Figma", level=2)
    add_bullets(doc, [
        "Use role-based navigation with clear labels: AP Board, Head Board, Invoices, Operations, Dashboard.",
        "Use consistent badges for AI Ready, Needs Review, High Risk, Missing Docs, Approval Pending and Payment Complete.",
        "Use icon language from the swimlane: human, AI agent, assist and mandatory approval.",
        "Cards should be dense and operational, not marketing-style; finance users need scanability.",
        "Every AI suggestion must show confidence, reason and the human action required.",
        "Critical buttons must match role permissions: approve/reject only for head, sign only for CFO, final financial edits only for AP/admin.",
    ])

    doc.add_heading("10. Notifications, Reporting and Monitoring", level=1)
    add_bullets(doc, [
        "Daily AP digest: pending approvals, missing docs, high-risk tickets, Xero failures and payment queue.",
        "Overdue workflow monitor: alerts when due date or SLA is near breach.",
        "Self-healing monitor: retry safe Xero token refresh, stuck agent jobs and transient gateway failures.",
        "Exception report: duplicate risk, account mismatch, low confidence OCR, payment confirmation mismatch.",
        "Performance dashboard: manual touches removed, average cycle time, first-pass approval rate, agent confidence and override rate.",
    ])

    doc.add_heading("11. Security, Compliance and Governance", level=1)
    add_bullets(doc, [
        "Agents must call existing APIs and respect RBAC; no direct database mutation for critical status/payment actions.",
        "Use prompt-injection defenses: document text is untrusted input and cannot instruct agents to bypass workflow rules.",
        "Encrypt uploaded documents and store least-privilege credentials for Xero/payment gateway.",
        "Mask sensitive bank account data in prompts where full value is not required.",
        "Log all agent actions with input reference, model/tool version, confidence and human acceptance.",
        "Require dual control for vendor bank account changes and payment release above threshold.",
    ])

    doc.add_heading("12. Testing Strategy", level=1)
    add_table(doc, ["Test Area", "Coverage"], [
        ("Unit tests", "Due-date cutoff, WHT calculation, validation rules, idempotency and risk scoring."),
        ("API tests", "Agent job creation, recommendation acceptance/rejection, RBAC and audit records."),
        ("Workflow tests", "Department upload to head approval, AP review, Xero, payment gateway and close."),
        ("AI evaluation tests", "OCR accuracy, duplicate detection, vendor/account matching and missing-doc classification."),
        ("Security tests", "Prompt injection documents, unauthorized role transitions, payment duplication and data leakage."),
        ("Regression tests", "Existing invoice/PO/ticket workflows continue working without agents enabled."),
    ])

    doc.add_heading("13. Phased Implementation Plan", level=1)
    add_table(doc, ["Phase", "Scope", "Business Value"], [
        ("Phase 1: Assistive AI", "Document extraction, validation checklist, missing-doc suggestions, AP summary.", "Reduces data entry and review time quickly."),
        ("Phase 2: Semi-autonomous workflow", "Agent jobs, recommendations, reminders, duplicate risk, WHT/voucher draft.", "Cuts repetitive AP work while humans approve key decisions."),
        ("Phase 3: Integration agents", "Xero assistant, payment gateway reconciliation, notification automation.", "Reduces finance close and reconciliation effort."),
        ("Phase 4: Predictive operations", "SLA prediction, anomaly detection, self-healing jobs, conversational assistant.", "Improves control, forecasting and exception handling."),
    ])
    doc.add_heading("Implementation Epics", level=2)
    add_table(doc, ["Epic", "Includes", "Primary Screens / Services"], [
        ("Agent foundation", "AgentJob, AgentRun, AgentRecommendation, event triggers, retry/idempotency.", "API services, Operations Agent Monitor."),
        ("Intake intelligence", "OCR extraction, confidence, field review, invoice/PO sync checks.", "Department Intake, Invoice Detail."),
        ("Validation intelligence", "Duplicate, missing docs, vendor/account, risk score, readiness state.", "Invoice Detail, Ticket Detail, AP Board badges."),
        ("Approval intelligence", "Head summary, rejection reason, reminders, SLA escalation.", "Head Board, Notifications."),
        ("AP copilot", "Next action, WHT/voucher draft, GL/expense suggestion, Xero payload draft.", "Ticket Detail."),
        ("Payment intelligence", "Payment batch validation, gateway confirmation, reconciliation.", "CFO View, Ticket Detail, Operations."),
        ("Monitoring and QA", "Workflow validation agent, failed job monitor, dashboard metrics.", "Dashboard, Operations, QA reports."),
    ])
    doc.add_heading("Release Acceptance Criteria", level=2)
    add_bullets(doc, [
        "All human approval gates remain enforced by backend role/status validation.",
        "No agent can directly execute high-risk payment or vendor bank account change.",
        "Every agent recommendation has confidence, reason, timestamp and audit trail.",
        "Clickable demo covers happy path, rejection, missing docs, high-risk vendor, CFO sign and Xero failure.",
        "Workflow QA agent can run at least one end-to-end scenario without manual test setup.",
    ])

    doc.add_heading("14. Immediate Implementation Backlog", level=1)
    add_numbered(doc, [
        "Create AgentJob, AgentRun and AgentRecommendation tables with audit-friendly metadata.",
        "Add background worker/orchestrator service and event triggers for invoice upload, approval, AP stage, Xero and payment events.",
        "Build Document Intake Agent for OCR/extraction with confidence fields.",
        "Build Validation/Risk Agent for duplicate, missing docs, PO sync and vendor account checks.",
        "Add AI extraction and validation UI cards to Invoice Detail and Ticket Detail.",
        "Add agent badges to AP Board and Head Board.",
        "Add Notification Agent for reminders and missing-doc follow-ups.",
        "Add QA/Workflow Agent with Playwright/API scenario tests.",
    ])

    doc.add_heading("15. Success Metrics", level=1)
    add_bullets(doc, [
        "Reduce clean invoice processing from roughly 3 days toward 1 day.",
        "Reduce manual data entry by 70-85%.",
        "Reduce missing-doc follow-up effort by 60-80%.",
        "Reduce daily reporting/follow-up effort by 80-90%.",
        "Maintain zero unauthorized payment release and zero duplicate Xero bill creation.",
        "Track AI override rate and keep high-risk auto-actions at 0 without human approval.",
    ])

    path = OUT / "AP_Automation_Agentic_AI_Detailed_PDR.docx"
    return save_doc(doc, path)


def master_doc(diagrams):
    doc = Document()
    style_doc(doc, "Company AP Automation Documentation Pack", "Business canvas, process flow, architecture, integrations, swimlane, application design and PRD")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("This documentation pack reflects the current project implementation in the local repository. The solution consolidates department invoice creation, PO synchronization, department head approval, AP finance processing, Xero bookkeeping, bank execution and requester closeout into one role-based AP automation application.")
    doc.add_heading("Implemented Scope Snapshot", level=1)
    add_bullets(doc, [
        "React frontend with login, AP/head board, dashboard, operations, invoice detail and ticket detail screens.",
        "NestJS API with auth, invoices, approvals, tickets, payment operations, Xero and reference modules.",
        "Prisma/Postgres domain model covering invoices, POs, tickets, activities, attachments, vendors, payments, approvals and notifications.",
        "Department head approval column and read-only decision portal.",
        "Comments, downloadable attachments and closed-ticket audit locking.",
        "Payment gateway flow for bank execution, Xero paid marking, requester notification and ticket close.",
    ])
    doc.add_heading("1. Business Canvas", level=1)
    add_image(doc, diagrams["canvas"], "Business canvas")
    doc.add_heading("2. Process Flow", level=1)
    add_image(doc, diagrams["process"], "Process flow")
    doc.add_heading("3. Functional Architecture Diagram", level=1)
    add_image(doc, diagrams["architecture"], "Functional architecture")
    doc.add_heading("4. Integration Diagram", level=1)
    add_image(doc, diagrams["integration"], "Integration diagram")
    doc.add_heading("5. Swimlane Diagram", level=1)
    add_image(doc, diagrams["swimlane"], "Swimlane diagram")
    doc.add_heading("6. Application Design", level=1)
    add_image(doc, diagrams["app"], "Application design")
    add_table(doc, ["Role", "Responsibility"], roles)
    doc.add_heading("7. Product Requirements Document", level=1)
    doc.add_paragraph("The PRD defines the product goal as replacing fragmented AP tracking with a governed, role-scoped workflow that reduces cycle time and manual entry while preserving finance controls.")
    add_table(doc, ["Step", "Stage", "Owner", "Output"], workflow_steps)
    doc.add_heading("Integration Notes", level=1)
    add_bullets(doc, [
        "Xero: real OAuth and Accounting API code exists for bill creation and payment marking when tenant credentials are configured.",
        "Bank/payment gateway: current implementation supports Meezan-style CSV export/import concepts and payment gateway execution flow. Production should use bank API/file response integration.",
        "Legacy data: schema and operations include old sheet references and Google CSV import/migration support.",
        "Trello: ticket fields store Trello card IDs/URLs for continuity; full Trello API sync is future scope.",
    ])
    doc.add_heading("Traceability", level=1)
    add_table(doc, ["Area", "Current Implementation"], [
        ("Workflow status engine", "api/src/tickets/tickets.service.ts"),
        ("Approval logic", "api/src/approvals and api/src/invoices"),
        ("Xero integration", "api/src/full-scope/full-scope.service.ts"),
        ("Payment gateway flow", "api/src/tickets/tickets.service.ts and dev-mock-api.cjs"),
        ("Role UI", "web/src/pages and web/src/components/Layout.tsx"),
    ])
    path = OUT / "AP_Automation_Master_Pack.docx"
    return save_doc(doc, path)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "canvas": business_canvas(),
        "process": process_flow(),
        "architecture": architecture_diagram(),
        "integration": integration_diagram(),
        "swimlane": swimlane(),
        "app": app_design_diagram(),
        "old_new": old_vs_new_artifacts_diagram(),
        "human_agent": human_agent_swimlane_diagram(),
    }
    docs = [
        business_canvas_doc(diagrams),
        application_design_doc(diagrams),
        prd_doc(diagrams),
        agentic_pdr_doc(diagrams),
        master_doc(diagrams),
    ]
    summary = OUT / "DELIVERABLES.txt"
    summary.write_text(
        "\n".join(
            ["AP Automation documentation deliverables", "", "JPG diagrams:"]
            + [f"- {p.name}" for p in diagrams.values()]
            + ["", "Word documents:"]
            + [f"- {p.name}" for p in docs]
        ),
        encoding="utf-8",
    )
    print(summary)


if __name__ == "__main__":
    main()
