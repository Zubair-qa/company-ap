from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "deliverables" / "ap_automation_prd"
DOCX_PATH = OUT_DIR / "AP_Automation_PRD_New_Scope.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "64748B"
LINE = "D7DEE8"
SOFT = "F2F4F7"
TEAL_FILL = "ECFDF5"
BLUE_FILL = "EFF6FF"
AMBER_FILL = "FFF7ED"
RED_FILL = "FEF2F2"
GREEN_FILL = "F0FDF4"
WHITE = "FFFFFF"
BLACK = "111827"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run(run, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_para_format(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, color=BLACK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_format(p, after=0, line=1.05)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("Company AP Automation PRD | New Scope")
    set_run(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Confidential working product requirements document")
    set_run(run, size=8.5, color=MUTED)


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    set_para_format(p, before=18, after=4)
    r = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run(r, size=12, bold=True, color=MUTED)

    p = doc.add_paragraph()
    set_para_format(p, after=6)
    r = p.add_run("Accounts Payable Automation Platform")
    set_run(r, size=26, bold=True, color=NAVY)

    p = doc.add_paragraph()
    set_para_format(p, after=18)
    r = p.add_run("Department invoice intake, AI validation, AP finance execution, CFO payment sign, payment gateway/bank execution, Xero bookkeeping, notification, and closeout.")
    set_run(r, size=13, color=MUTED)

    meta = [
        ("Version", "1.0 - New Scope Baseline"),
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Primary workflow", "Department User -> AI/Validation -> AP Finance -> CFO -> Payment Gateway/Bank -> Xero/Notify/Close"),
        ("Explicitly removed", "Separate Department Head dashboard/workflow. Department approval is now a required document attachment."),
        ("Prepared for", "AP Automation implementation and Figma clickable demo alignment"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1800, 7560])
    for i, (label, value) in enumerate(meta):
        shade_cell(table.cell(i, 0), SOFT)
        set_cell_text(table.cell(i, 0), label, bold=True, color=NAVY)
        set_cell_text(table.cell(i, 1), value)

    add_callout(
        doc,
        "Executive intent",
        "The product consolidates Trello, Google Sheets, Excel trackers, email invoices, manual vouchers, bank portal steps, and Xero bookkeeping into a single controlled AP workspace. The target business outcome is to reduce a roughly three-day manual cycle toward a one-day clean-invoice cycle while lowering duplicate entry, missing-document loops, payment risk, and audit gaps.",
        BLUE_FILL,
        BLUE,
    )
    doc.add_page_break()


def add_callout(doc: Document, title: str, body: str, fill: str, border: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_format(p, after=4)
    r = p.add_run(title)
    set_run(r, size=11, bold=True, color=border)
    p2 = cell.add_paragraph()
    set_para_format(p2, after=0, line=1.10)
    r2 = p2.add_run(body)
    set_run(r2, size=10, color=BLACK)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    set_para_format(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para_format(p, after=4, line=1.10)
        run = p.add_run(item)
        set_run(run, size=10.8)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para_format(p, after=4, line=1.10)
        run = p.add_run(item)
        set_run(run, size=10.8)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=SOFT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        shade_cell(hdr[idx], header_fill)
        set_cell_text(hdr[idx], header, bold=True, color=NAVY, size=9.2)
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row[idx], value, size=9)
    doc.add_paragraph()
    return table


def page(doc: Document, title: str):
    if len(doc.paragraphs) > 1:
        doc.add_page_break()
    add_heading(doc, title, 1)


def add_diagram(doc: Document, path: Path, caption: str):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.2))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(cp, after=10)
        r = cp.add_run(caption)
        set_run(r, size=9, italic=True, color=MUTED)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    page(doc, "1. Executive Summary and Product Vision")
    add_para(doc, "The Accounts Payable Automation Platform is a role-based finance operations system that replaces fragmented AP handling across Trello, Google Sheets, Excel files, email, manual bank portal work, and separate bookkeeping updates. The new scope centers on a single operational chain: Department User -> AI/Validation -> AP Finance -> CFO -> Payment Gateway/Bank -> Xero/Notify/Close.")
    add_para(doc, "The product should make invoice requests traceable from the moment a department creates them until payment completion. Every invoice should own its synced PO, required document package, AI validation result, AP finance checks, WHT calculation, voucher, CFO sign event, payment gateway/bank status, Xero bill/payment record, requester notification, and audit trail.")
    add_callout(doc, "North-star outcome", "Move clean AP requests from a roughly three-day manual turnaround toward a one-day controlled workflow by removing duplicate data entry, preventing missing-document surprises, and automating low-risk repetitive tasks while preserving CFO payment authorization.", TEAL_FILL, "0F766E")
    add_heading(doc, "Objectives", 2)
    add_bullets(doc, [
        "Centralize all invoice and payment requests into one board and one data model.",
        "Ensure departments can register, sign in, create invoice requests, and view only their own records.",
        "Require a complete document package before AP Finance takes ownership.",
        "Use agentic AI to extract fields, validate documents, score confidence, and identify missing or risky items.",
        "Preserve mandatory human controls for AP exception handling, tax/account overrides, vendor bank changes, and CFO payment sign.",
        "Integrate Xero so bookkeeping records are created and payment status is reflected.",
        "Integrate a payment gateway/bank adapter so payment execution status can drive closeout.",
    ])

    page(doc, "2. Background, Current Problems, and Business Case")
    add_para(doc, "The existing AP process is distributed across multiple tools and people. Invoices may arrive through department uploads, email, scanned slips, repair/maintenance bills, Urdu-language payment notes, Google Sheets, Excel trackers, old reference sheets, and Trello cards. Finance then performs repeated checks for vendor account numbers, PO references, invoice number quality, WHT filer/non-filer status, voucher generation, bank upload, CFO sign, bank execution, Xero marking, and requester updates.")
    add_heading(doc, "Manual issues to eliminate", 2)
    add_bullets(doc, [
        "Duplicate entry across Trello, Google Sheets, Excel, Xero, and the bank portal.",
        "No single source of truth for ticket status, blocker, missing document, assignment, or owner.",
        "Late discovery of missing PO, invoice, slip, GR, department approval, or vendor bank proof.",
        "Invoice numbers are not always unique, and some slips do not contain invoice numbers.",
        "Vendor account numbers may be absent from invoice documents and require old-sheet verification.",
        "Partial payment flows require old-ticket and old-sheet references.",
        "Manual WHT calculation and voucher preparation create error risk.",
        "CFO sign and bank execution are tracked manually, making closeout delayed and hard to audit.",
    ])
    add_heading(doc, "Business value", 2)
    add_table(doc, ["Value Driver", "Current Pain", "Target Improvement"], [
        ["Cycle time", "Manual follow-up and duplicate entry stretch clean requests over multiple days.", "One-day clean-invoice target through AI validation, structured routing, and automated status updates."],
        ["Control", "Approvals and documents are scattered across email, sheets, and cards.", "Required documents are attached to the invoice ticket and validated before AP processing."],
        ["Accuracy", "Vendor, PO, account, tax, and payment status checks happen manually.", "AI pre-checks and AP review reduce preventable errors before payment."],
        ["Audit", "Status history and evidence are fragmented.", "Every action, attachment, integration event, and exception is logged against the ticket."],
    ], [1900, 3750, 3710])

    page(doc, "3. Scope, Non-Scope, and Product Boundaries")
    add_heading(doc, "In scope", 2)
    add_bullets(doc, [
        "Department registration and role-scoped sign-in.",
        "Department-owned invoice creation with source/type selection: slip/image, PDF/email invoice, manual entry, Excel import, and Google Sheet import.",
        "Synced invoice and purchase order creation from the same request context.",
        "Required document package validation: department approval document, PO, invoice/slip, GR or receiving proof, vendor bank proof, and other supporting documents as configured.",
        "OCR and document intelligence agent that extracts fields and computes verification percentage.",
        "AP Finance kanban board and ticket detail workspace.",
        "Vendor, PO, invoice account, old sheet reference, WHT, voucher, Xero, payment gateway/bank, notification, and closeout workflows.",
        "CFO sign queue with human approval before payment execution.",
        "Role-based editing, read-only closed tickets, comments, attachments, download, audit log, alerts, and status notifications.",
        "Legacy migration support for old Google Sheets, Excel trackers, Trello references, and old payment references.",
    ])
    add_heading(doc, "Out of scope for the new baseline", 2)
    add_bullets(doc, [
        "Separate Department Head dashboard, board, approval queue, or workflow stage.",
        "Department Head as a system role for live routing. Department approval is treated as a required attachment/evidence document.",
        "Stripe checkout as a payment flow. The target payment integration is payment gateway/bank adapter based.",
        "Manual AP-created tickets that are detached from invoices. Invoice creation should generate the AP ticket automatically.",
        "Fully autonomous payment release without CFO sign and configured payment controls.",
    ])
    add_callout(doc, "Implementation cleanup note", "The current repository still contains some legacy concepts such as DEPT_ADMIN and Department Head approval statuses in parts of the schema/UI. The target PRD scope requires removing or deprecating these from live workflow behavior while preserving historical audit/migration compatibility where needed.", RED_FILL, "B91C1C")

    page(doc, "4. Users, Roles, Access, and Permissions")
    add_para(doc, "The product must enforce access by role, department, ticket status, and field ownership. Users should never see or edit more than their workflow requires. All mutations should generate user-visible success/error alerts and an audit activity entry.")
    add_table(doc, ["Role", "Primary Responsibilities", "Access Boundary"], [
        ["Department User", "Register department, create invoices, attach required documents, fix missing documents, view own department board and ticket status.", "Own department only. Can edit draft/rework fields and attachments before AP ownership."],
        ["AP Clerk / AP Finance", "Review AI-validated tickets, verify documents, vendor, PO and account details, calculate WHT, generate voucher, prepare Xero/payment steps, request missing docs.", "Finance workspace. Cannot CFO-sign or bypass required document controls."],
        ["CFO", "Review payment summary, risk indicators, AP verification, and sign payment before execution.", "CFO sign queue and payment summary. Cannot alter department invoice evidence or AP calculations except returning for correction."],
        ["Company Admin", "Configure departments, vendors, roles, reference data, agent settings, integrations, reporting, and migration jobs.", "All departments and admin views, with sensitive changes audited."],
        ["System / AI Agents", "Extract, validate, score, recommend, reconcile, notify, monitor, and escalate.", "No final approval, tax override, vendor bank change, or payment release authority."],
    ], [1500, 4700, 3160])
    add_heading(doc, "Permission principles", 2)
    add_bullets(doc, [
        "Department users can create and correct requests but cannot process WHT, voucher, Xero, or payment execution.",
        "AP Finance can process operational finance fields but cannot sign payments as CFO.",
        "CFO can sign or return payment but should not edit underlying invoice/PO document fields.",
        "Closed payment tickets are read-only for all regular roles and require a linked follow-up ticket for remaining/partial payment.",
        "Vendor bank account changes require explicit human approval and audit logging.",
    ])

    page(doc, "5. End-to-End Workflow")
    add_diagram(doc, ROOT / "deliverables" / "ap_automation_docs" / "05_swimlane_diagram.jpg", "Figure 1. Updated role swimlane without separate Department Head workflow.")
    add_numbered(doc, [
        "Department user signs in or registers a new department and creates an invoice request.",
        "The request captures invoice fields, PO fields, vendor, amount, expense nature, payment method, and required documents.",
        "The AI/OCR agent extracts field values from uploaded slip, image, PDF, email invoice, or imported sheet row.",
        "The validation agent checks required documents, PO sync, vendor match, account proof, duplicate invoices, partial payment references, and confidence.",
        "If validation is incomplete, the ticket stays in Draft/Rework and shows missing documents and AI recommendations to the department.",
        "If validation passes configured thresholds, the invoice-generated ticket moves to AP Finance.",
        "AP Finance performs document review, vendor/PO/account verification, WHT calculation, voucher generation, and Xero bill preparation.",
        "AP prepares bank/payment gateway upload and moves the ticket to CFO sign pending.",
        "CFO verifies the payment summary and signs or returns the payment.",
        "Payment gateway/bank executes the payment and returns confirmation.",
        "Xero is marked paid, requester is notified, and the ticket is closed/read-only.",
    ])
    add_heading(doc, "Due date rule", 2)
    add_para(doc, "Due date calculation must use a configurable daily cutoff time in the company timezone. If the request reaches Finance before the configured cutoff time, the due date is counted from the same business date plus three days. If it reaches Finance after the configured cutoff time, the due date is counted from the next business date plus three days. The cutoff time must be configurable rather than hard-coded.")

    page(doc, "6. Invoice Intake, Department Registration, and Document Package")
    add_heading(doc, "Department registration", 2)
    add_para(doc, "The sign-in screen must include a registration option so a new department can be created without manually editing database seed data. Registration should capture department name, department code, requester/admin email, active status, and the named owner or source for department approval documents. Registration should create a pending admin-review record unless the company permits automatic activation.")
    add_heading(doc, "Invoice source/type handling", 2)
    add_table(doc, ["Source Type", "Form Behavior", "AI Behavior"], [
        ["Slip/image", "Show upload first, then auto-populated fields for review.", "OCR classifies bill type, vendor, date, amount, account number, PO hints, and confidence."],
        ["PDF/email invoice", "Show attachment upload and extracted invoice fields.", "Extract invoice number, vendor, amount, tax, terms, bank details, and PO/reference data."],
        ["Manual entry", "Show full invoice and PO fields because no extraction source exists.", "Validate entered fields, duplicates, document presence, and reference data."],
        ["Excel/Google Sheet", "Hide manual invoice fields and show mapping/import controls.", "Map columns, flag invalid rows, create invoice tickets in batch, and retain legacy row reference."],
    ], [1800, 3600, 3960])
    add_heading(doc, "Required document package", 2)
    add_bullets(doc, [
        "Department approval document or email approval converted/attached as evidence.",
        "Purchase order, with PO number and vendor/amount sync.",
        "Supplier invoice, scanned bill, repair/maintenance slip, email invoice, or equivalent source.",
        "GR, GRN, receiving proof, delivery note, or service completion proof where applicable.",
        "Vendor bank proof, account confirmation, or old-sheet verification reference when invoice lacks account number.",
        "Voucher, bank confirmation, and Xero records generated later in the workflow.",
    ])

    page(doc, "7. Agentic AI Validation Requirements")
    add_para(doc, "The AI layer should operate as a controlled set of agents that perform extraction, validation, recommendation, and monitoring without bypassing human gates. The system must store the agent output, confidence score, evidence references, and any human override.")
    add_table(doc, ["Agent", "Responsibilities", "Output"], [
        ["OCR / Document Intake Agent", "Reads images, PDFs, email invoices, slips, and imported rows. Classifies bill type and extracts invoice fields.", "Extracted JSON, field-level confidence, source page/region references."],
        ["Document Completeness Agent", "Checks required documents by bill type and department policy.", "Verification percentage, missing document list, pass/fail gate."],
        ["Vendor / PO / Account Agent", "Matches vendor, PO, invoice account and old-sheet account references.", "Match status, mismatch reason, manual review recommendation."],
        ["Duplicate / Partial Payment Agent", "Detects duplicate invoice numbers, same vendor/amount/date patterns, and 50/50 partial-payment relationships.", "Duplicate risk score, parent ticket reference, old sheet reference."],
        ["WHT / Voucher Assistant", "Suggests filer/non-filer status, WHT rate, WHT amount, net payable, and voucher draft.", "Recommended calculation and voucher draft for AP approval."],
        ["Reconciliation / Notification Agent", "Reconciles bank confirmation, Xero payment status, ticket status, and requester notification.", "Close recommendation, mismatch alert, notification event."],
    ], [2050, 4800, 2510])
    add_heading(doc, "AI verification percentage", 2)
    add_para(doc, "The UI must show a clear percentage such as 82% verified. The percentage should be derived from document completeness, field confidence, PO/vendor/account match, duplicate risk, and policy-critical checks. Low confidence should not silently pass; it must create a visible missing-item or review recommendation.")
    add_callout(doc, "Human-in-the-loop boundary", "AI may recommend, pre-fill, summarize, reconcile, and notify. AI must not approve spend, sign payments, override tax, change vendor bank details, or release funds without an authorized human action.", AMBER_FILL, "B45309")

    page(doc, "8. AP Finance Processing Requirements")
    add_para(doc, "AP Finance receives only invoice-generated tickets that have passed AI validation or were explicitly sent as exception cases. AP must have a kanban board and ticket detail screen that make status, assigned user, due date, priority, missing documents, document status, Xero status, bank status, and WHT status immediately visible.")
    add_heading(doc, "AP workflow states", 2)
    add_bullets(doc, [
        "AI validated documents / Docs review",
        "Vendor, PO, and account verification",
        "WHT calculation",
        "Voucher generation",
        "Xero bill entry",
        "Payment preparation",
        "Bank upload",
        "CFO sign pending",
        "Bank execution pending",
        "Bank executed",
        "Marked paid in Xero",
        "Requester notified",
        "Payment complete",
    ])
    add_heading(doc, "AP controls", 2)
    add_bullets(doc, [
        "AP can request documents back from the department with reason and comments.",
        "AP can assign tickets within the finance team based on scope and role.",
        "AP can accept or adjust AI WHT/voucher recommendations, with calculation audit retained.",
        "AP can prepare but not CFO-sign payment.",
        "AP can close only after payment gateway/bank execution and Xero paid status are recorded or reconciled.",
    ])

    page(doc, "9. CFO, Payment Gateway/Bank, Xero, Notify, and Close")
    add_heading(doc, "CFO sign", 2)
    add_para(doc, "The CFO portal/queue must show only sign-pending items and the supporting risk summary. CFO must see vendor, amount, net payable, WHT, payment method, bank account status, document package score, AP verifier, Xero bill reference, and any high-risk flags. CFO can sign or return with a reason.")
    add_heading(doc, "Payment gateway/bank execution", 2)
    add_para(doc, "After CFO sign, payment execution moves to the payment gateway/bank adapter. The adapter should support a test/sandbox mode for implementation and a production bank/API mode later. Execution results must update bank status, payment reference, executed timestamp, and reconciliation status.")
    add_heading(doc, "Xero bookkeeping", 2)
    add_bullets(doc, [
        "Create or match Xero contact from vendor master.",
        "Create Xero bill after AP verification and before payment close.",
        "Store Xero bill ID, bill number, contact ID, sync status, last sync time, and error messages.",
        "Mark payment as paid in Xero only after bank/payment gateway execution confirmation.",
        "Retry failed syncs safely and surface failures in the agent monitor.",
    ])
    add_heading(doc, "Notify and close", 2)
    add_para(doc, "On successful reconciliation, the system notifies the requester/department, updates the board, closes the ticket, and locks it as read-only. Any follow-up or remaining partial payment should create a linked child ticket rather than reopening the closed record.")

    page(doc, "10. Functional Requirements")
    add_table(doc, ["ID", "Requirement", "Acceptance Criteria"], [
        ["FR-001", "Department registration is available from sign-in.", "A new department request can be submitted with name, code, requester email, and approval document owner."],
        ["FR-002", "Role-scoped login and navigation.", "Each role sees only its allowed dashboards, boards, and actions after login."],
        ["FR-003", "Department invoice creation creates a synced invoice/PO/ticket context.", "No AP-only detached manual ticket is required for standard invoice flow."],
        ["FR-004", "Source/type dropdown changes form behavior.", "Manual mode shows fields; sheet import hides manual invoice fields; upload modes show OCR extraction."],
        ["FR-005", "Required documents are enforced.", "Missing department approval, PO, invoice/slip, GR, or bank proof prevents clean AP routing."],
        ["FR-006", "AI validation score is visible.", "User sees percentage verified, missing items, confidence, and next action."],
        ["FR-007", "AP kanban follows scoped statuses.", "Cards move from AI validated docs through payment complete with role-based transitions."],
        ["FR-008", "CFO sign is mandatory before payment execution.", "Payment gateway/bank execution cannot occur without CFO sign event."],
        ["FR-009", "Xero sync records bill/payment status.", "Bill created and paid-marked events are persisted with IDs and errors."],
        ["FR-010", "Closed tickets are locked.", "Payment complete tickets cannot be edited except via admin/audit process or linked follow-up."],
    ], [950, 4050, 4360])

    page(doc, "11. Non-Functional Requirements")
    add_table(doc, ["Category", "Requirement"], [
        ["Security", "JWT authentication, role guards, department scoping, least-privilege API permissions, encrypted secrets, and audit logs for sensitive actions."],
        ["Performance", "Board and ticket APIs should load within acceptable interactive thresholds for daily AP operations; background agents should not block page responsiveness."],
        ["Reliability", "Integrations with Xero and payment gateway/bank must be retryable, idempotent where possible, and visible in an agent monitor."],
        ["Auditability", "Every status change, attachment, comment, AI result, integration call, override, and payment action must be traceable."],
        ["Usability", "Users should see clear success/error alerts after saving, uploading, submitting, signing, syncing, or closing."],
        ["Data quality", "Invoice numbers may not be unique globally; uniqueness must rely on internal reference plus vendor/date/amount matching and duplicate-risk logic."],
        ["Accessibility", "Forms, boards, buttons, validation messages, and document statuses should be keyboard-accessible and readable."],
        ["Configuration", "Cutoff time, required documents by bill type, WHT rates, gateway settings, Xero credentials, and role permissions must be configurable."],
    ], [1900, 7460])

    page(doc, "12. Data Model and Core Artifacts")
    add_para(doc, "The current repository uses Prisma with PostgreSQL and already models many important AP entities. The target PRD keeps these concepts but updates the workflow scope to remove live Department Head routing.")
    add_table(doc, ["Artifact / Entity", "Purpose", "Key Fields"], [
        ["Department", "Department ownership and scoping.", "Code, name, active status, users."],
        ["User", "Authentication and role-based ownership.", "Email, password hash, role, department, active, last login."],
        ["Vendor", "Vendor master and payment data.", "Display/legal name, tax IDs, bank details, terms, WHT rate."],
        ["Invoice", "Department-created invoice record.", "Internal ref, invoice number, amount, dates, extracted JSON, status, department, vendor, PO."],
        ["PurchaseOrder", "Synced PO context.", "PO number, vendor, department, requester, totals, status, line items."],
        ["PaymentTicket", "Operational AP workflow card.", "Status, priority, assignee, due date, document status, missing docs, WHT, Xero, bank status."],
        ["SupportingDocument", "Attached evidence and downloads.", "Document type, file name/path, MIME type, size, uploader, invoice/ticket/PO link."],
        ["TicketActivity", "Audit and comments.", "Actor, message, from/to status, timestamp."],
        ["PaymentRecord", "Payment schedule/execution/reconciliation.", "Vendor, amount, status, reference, authorized by, reconciliation."],
        ["XeroConnection", "Accounting integration credentials/status.", "Tenant, token metadata, connected/synced state."],
    ], [1700, 3000, 4660])
    add_heading(doc, "Required project artifacts", 2)
    add_bullets(doc, [
        "Business canvas",
        "Process flow diagram",
        "Functional architecture diagram",
        "Integration diagram",
        "Human/AI swimlane diagram",
        "Application design / clickable Figma prototype",
        "Old vs new artifacts comparison",
        "Product Requirements Document",
        "Data dictionary and migration mapping",
        "API contract and integration specs",
        "QA test plan and UAT scripts",
        "Security and audit matrix",
    ])

    page(doc, "13. Technical Architecture and Tech Stack")
    add_heading(doc, "Current implementation stack", 2)
    add_table(doc, ["Layer", "Technology"], [
        ["Frontend", "React 19, TypeScript, Vite, React Router, TanStack React Query, Axios."],
        ["Backend", "NestJS 11, TypeScript, Passport/JWT, class-validator, class-transformer, Multer."],
        ["Database / ORM", "PostgreSQL via Prisma 5."],
        ["File handling", "Multer-backed uploads with SupportingDocument records and downloadable attachments."],
        ["Spreadsheet handling", "xlsx package for Excel/CSV import and old sheet migration flows."],
        ["Accounting", "Xero integration module/adapter with OAuth/API credential configuration."],
        ["Payment", "Payment gateway/bank adapter. Stripe checkout is not part of target scope."],
        ["AI", "Planned agentic AI service layer for OCR, validation, risk scoring, WHT suggestions, reconciliation, and monitoring."],
        ["Testing", "Jest, Nest testing, frontend TypeScript build/lint. Add workflow E2E tests for scoped routes."],
    ], [2100, 7260])
    add_diagram(doc, ROOT / "deliverables" / "ap_automation_docs" / "03_functional_architecture.jpg", "Figure 2. Functional architecture artifact, to be kept aligned with payment gateway/bank target scope.")
    add_heading(doc, "Recommended production architecture", 2)
    add_bullets(doc, [
        "Frontend app served via CDN/static hosting with environment-specific API base URL.",
        "NestJS API behind HTTPS with JWT authentication, request logging, and role guards.",
        "PostgreSQL managed database with backups, migrations, and row-level scoping at application layer.",
        "Object/file storage for invoices and supporting documents.",
        "Background job queue for OCR, AI validation, Xero sync, payment status polling, notifications, and retries.",
        "Secrets manager for Xero, payment gateway/bank, OCR/LLM, email, and database credentials.",
        "Observability stack for logs, traces, metrics, failed agent runs, and integration retries.",
    ])

    page(doc, "14. Integrations and Migration")
    add_heading(doc, "Integration scope", 2)
    add_table(doc, ["System", "Purpose", "Behavior"], [
        ["Xero", "Bookkeeping and payment record sync.", "Create bill, update paid status, store IDs/errors, retry failed syncs."],
        ["Payment Gateway / Bank", "Payment upload, sign handoff, execution confirmation.", "Adapter supports sandbox/test mode and production status callbacks/polling."],
        ["Google Sheets / Excel", "Legacy import and batch invoice creation.", "Map columns, validate rows, retain row IDs and old references."],
        ["Trello", "Legacy reference migration and optional notification/update.", "Store old card ID/URL; future integration can backfill or close legacy cards."],
        ["Email / Notifications", "Requester and AP updates.", "Send missing-doc, overdue, payment-complete and integration-failure notifications."],
        ["OCR / LLM", "Document extraction and agentic validation.", "Return structured JSON, confidence, evidence, missing-item list, and recommendation."],
    ], [1650, 2700, 5010])
    add_heading(doc, "Migration requirements", 2)
    add_bullets(doc, [
        "Import old Google Sheet v1/v2 and Excel rows into normalized invoice/ticket records.",
        "Retain legacy row ID, sheet name, Trello card ID/URL, old reference, and parent/child ticket link for partial payments.",
        "Deduplicate vendors and map old vendor names to vendor master records.",
        "Preserve paid history and do not reopen old completed tickets unless a linked follow-up payment is created.",
        "Create reconciliation reports comparing imported totals, paid totals, open totals, and exceptions.",
    ])

    page(doc, "15. Reporting, Dashboards, and Notifications")
    add_para(doc, "The reporting layer must serve daily AP operations and management visibility. Reports should be role scoped: departments see their own requests; AP sees operational workload; CFO sees sign queue and payment risk; company admin sees cross-department performance.")
    add_heading(doc, "Dashboards", 2)
    add_bullets(doc, [
        "Department board: draft/rework, AI validation, AP processing, paid/closed.",
        "AP board: validated docs, vendor/PO/account, WHT/voucher, CFO sign pending, bank execution, Xero paid, closed.",
        "CFO queue: sign pending, returned, signed, high-risk payments.",
        "Agent monitor: OCR jobs, validation failures, Xero sync retries, payment gateway failures, notification failures.",
        "Reports dashboard: cycle time, missing-doc loops, WHT totals, payment status, overdue tickets, manual effort reduced, agent confidence distribution.",
    ])
    add_heading(doc, "Notifications and alerts", 2)
    add_table(doc, ["Trigger", "Recipient", "Message"], [
        ["Missing required document", "Department user", "List missing documents and direct link to upload/fix screen."],
        ["AI validation passed", "AP Finance", "Ticket is ready for AP review with score and risk summary."],
        ["Overdue or due soon", "Ticket owner / AP lead", "Due date alert based on configured cutoff time and SLA."],
        ["CFO sign pending", "CFO", "Payment summary, amount, vendor, risk score, and sign action link."],
        ["Payment executed", "Requester and AP", "Payment reference, Xero status, and closeout status."],
        ["Integration failure", "Company admin / AP lead", "Failed system, error reason, retry action, and affected ticket."],
    ], [2350, 2350, 4660])

    page(doc, "16. Testing, QA, and Acceptance Plan")
    add_heading(doc, "Testing layers", 2)
    add_bullets(doc, [
        "Unit tests for due date rule, status transitions, role permissions, WHT calculation, and validation scoring.",
        "API tests for auth, department scoping, invoice creation, document upload/download, ticket transitions, Xero/payment actions, and comments.",
        "Frontend tests for role navigation, invoice source dropdown behavior, document checklist, AI score display, AP board, CFO sign, alerts, and read-only closed ticket behavior.",
        "Integration tests with Xero sandbox/mock and payment gateway/bank sandbox/mock.",
        "Migration tests using representative old Google Sheet, Excel, and Trello reference data.",
        "Security tests for cross-department access, unauthorized status moves, attachment download authorization, and sensitive action auditing.",
        "User acceptance tests for Department User, AP Clerk, CFO, and Company Admin journeys.",
    ])
    add_heading(doc, "Sample UAT scenarios", 2)
    add_numbered(doc, [
        "Register Procurement department, sign in as Procurement user, and confirm only Procurement tickets show.",
        "Create slip-based invoice with missing GR and verify AI percentage shows incomplete and blocks AP routing.",
        "Attach missing GR and bank proof, rerun validation, and confirm ticket moves to AP Finance.",
        "AP verifies vendor/PO/account, accepts WHT suggestion, generates voucher, and prepares payment.",
        "CFO signs payment and confirms AP/bank execution cannot occur before sign.",
        "Payment gateway/bank executes payment, Xero marks paid, requester is notified, and ticket becomes read-only.",
        "Create a remaining 50% partial payment linked to old paid ticket and verify old reference is retained.",
    ])

    page(doc, "17. Implementation Roadmap")
    add_table(doc, ["Phase", "Scope", "Exit Criteria"], [
        ["Phase 1 - Scope cleanup", "Remove Department Head workflow from active UI/API behavior, align statuses, update role permissions, remove Stripe checkout UX.", "New flow is enforced end to end in design and app navigation."],
        ["Phase 2 - Department intake", "Department registration, procurement department, source/type forms, invoice/PO sync, required document uploads/downloads.", "Department users can create and correct scoped requests."],
        ["Phase 3 - AI validation", "OCR/extraction, document completeness, verification percentage, missing-item logic, duplicate and PO/vendor/account checks.", "Requests cannot move to AP without configured validation or explicit AP exception."],
        ["Phase 4 - AP finance execution", "AP board, WHT, voucher, assignments, comments, alerts, Xero bill creation.", "AP can process validated tickets through CFO sign pending."],
        ["Phase 5 - Payment and close", "CFO sign, payment gateway/bank adapter, bank execution, Xero paid mark, notify/close, read-only lock.", "Successful payment flow closes automatically with audit history."],
        ["Phase 6 - Migration and reporting", "Legacy Google Sheet/Excel/Trello migration, old references, dashboards, agent monitor.", "Management can track cycle time, blockers, exceptions, and migration reconciliation."],
    ], [1500, 5000, 2860])
    add_heading(doc, "Definition of done", 2)
    add_bullets(doc, [
        "All core roles can complete their scoped happy paths.",
        "All status transitions are server-validated by role and current status.",
        "All save/update actions show success/error alerts and create activity logs.",
        "Attachments can be uploaded, previewed/listed, downloaded, and permission-checked.",
        "AI validation score and missing documents are visible and persisted.",
        "Xero and payment gateway/bank adapter events are idempotent and auditable.",
        "Closed payment tickets are locked and can only be followed by linked tickets.",
    ])

    page(doc, "18. Risks, Controls, and Open Decisions")
    add_table(doc, ["Risk / Decision", "Impact", "Recommended Control"], [
        ["AI extraction error", "Wrong vendor, amount, account, or PO data can delay or misroute payment.", "Field-level confidence, source evidence, AP review, and no autonomous payment release."],
        ["Missing document bypass", "AP processes incomplete requests.", "Server-side required-document policy, UI checklist, audit trail, and AP exception reason."],
        ["Vendor bank fraud", "Payment to wrong account.", "Vendor master approval, bank proof verification, change audit, and CFO risk summary."],
        ["Xero sync failure", "Bookkeeping diverges from payment status.", "Retry queue, agent monitor, error visibility, and reconciliation report."],
        ["Payment gateway/bank failure", "Ticket closes without actual payment.", "Close only after execution confirmation or manually audited exception."],
        ["Legacy migration duplicates", "Old paid or partial invoices duplicate current open tickets.", "Deduplication rules, old reference preservation, and migration reconciliation."],
        ["Role creep", "Users edit fields outside responsibility.", "Backend role guards, field-level permissions, route protection, and test coverage."],
    ], [2200, 3200, 3960])
    add_heading(doc, "Open decisions", 2)
    add_bullets(doc, [
        "Which payment gateway/bank API will be the production target, and what sandbox is available?",
        "Which OCR/LLM provider should be used for document extraction and Urdu/repair slip handling?",
        "What is the configured daily cutoff time and business calendar?",
        "Which document types are mandatory for each bill type, department, and payment method?",
        "What approval evidence format is acceptable for the department approval document?",
        "What thresholds allow AI validation to auto-route to AP versus return to department rework?",
    ])

    page(doc, "19. Appendix A - New Scope Workflow Reference")
    add_diagram(doc, ROOT / "deliverables" / "ap_automation_docs" / "08_human_agent_swimlane.jpg", "Figure 3. Human and AI agent ownership under the scoped workflow.")
    add_para(doc, "The new baseline removes the separate Department Head board and treats department approval as evidence inside the invoice document package. The operational workflow is therefore shorter, easier to test, and easier for users to understand.")
    add_heading(doc, "Final scoped flow", 2)
    add_para(doc, "Department User -> AI/Validation -> AP Finance -> CFO -> Payment Gateway/Bank -> Xero/Notify/Close")
    add_heading(doc, "Primary implementation artifacts", 2)
    add_bullets(doc, [
        "Updated Figma prototype: deliverables/figma_design/index.html",
        "Updated Figma flow map: deliverables/figma_design/figma_flow_map.md",
        "Updated swimlane diagrams: deliverables/ap_automation_docs/05_swimlane_diagram.jpg and 08_human_agent_swimlane.jpg",
        "This PRD: deliverables/ap_automation_prd/AP_Automation_PRD_New_Scope.docx",
    ])

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
